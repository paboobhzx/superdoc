import importlib
import importlib.util
import io
import json
import os
import shutil
import sys
import types
import unittest
import zipfile
from pathlib import Path


def _install_handler_stubs():
    repo_root = Path(__file__).resolve().parents[1]
    if str(repo_root) not in sys.path:
        sys.path.insert(0, str(repo_root))

    dynamo = types.SimpleNamespace(
        update_job=lambda *args, **kwargs: None,
        mark_done=lambda *args, **kwargs: None,
        mark_failed=lambda *args, **kwargs: None,
    )
    s3 = types.SimpleNamespace(
        get_bytes=lambda key: b"",
        put_bytes=lambda key, data: None,
        make_output_key=lambda job_id, file_key, filename: f"outputs/{job_id}/{filename}",
    )
    logger = types.SimpleNamespace(get_logger=lambda name: types.SimpleNamespace(
        info=lambda *args, **kwargs: None,
        exception=lambda *args, **kwargs: None,
    ))
    limits = types.SimpleNamespace(
        assert_pdf_page_limit=lambda *args, **kwargs: None,
        count_pdf_pages=lambda *args, **kwargs: 1,
        page_limit_for_user=lambda *args, **kwargs: 20,
        storage_ttl_for_user=lambda *args, **kwargs: 43200,
    )
    sys.modules["dynamo"] = dynamo
    sys.modules["s3"] = s3
    sys.modules["logger"] = logger
    sys.modules["limits"] = limits
    return dynamo, s3


def _minimal_docx_bytes(text):
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w") as zf:
        zf.writestr("[Content_Types].xml", "<Types xmlns=\"http://schemas.openxmlformats.org/package/2006/content-types\"/>")
        zf.writestr("word/document.xml", f"<w:document>{text}</w:document>")
    return out.getvalue()


def _minimal_pdf_bytes(n_pages: int = 1) -> bytes:
    """Create a minimal valid PDF with n blank pages using pypdf."""
    from pypdf import PdfWriter
    writer = PdfWriter()
    for _ in range(n_pages):
        writer.add_blank_page(width=612, height=792)
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()


def _text_pdf_bytes(text: str = "sample words " * 30) -> bytes:
    import pymupdf

    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    out = doc.tobytes()
    doc.close()
    return out


def _import_operations():
    layer_path = Path(__file__).resolve().parents[1] / "layers" / "superdoc_utils"
    if str(layer_path) not in sys.path:
        sys.path.insert(0, str(layer_path))
    sys.modules.pop("operations", None)
    return importlib.import_module("operations")


class OperationCatalogTests(unittest.TestCase):
    def setUp(self):
        self.operations = _import_operations()

    def test_markdown_catalog_includes_convert_and_editor_for_text_types(self):
        for input_type in ("md", "markdown", "txt"):
            ops = {item["operation"]: item for item in self.operations.list_operations(input_type)}
            self.assertIn("markdown_convert", ops)
            self.assertIn("md_edit", ops)
            self.assertEqual(ops["md_edit"]["kind"], "client_editor")
            self.assertEqual(ops["md_edit"]["intent"], "edit")
            self.assertEqual(ops["md_edit"]["editor_route"], "/editor/markdown")
            self.assertEqual(ops["markdown_convert"]["targets"], ["pdf", "docx", "png", "jpg", "jpeg", "tiff", "md", "txt", "html"])

    def test_docx_catalog_includes_edit_text_and_pdf(self):
        ops = {item["operation"] for item in self.operations.list_operations("docx")}
        self.assertTrue({"doc_edit", "docx_to_txt", "docx_to_pdf"}.issubset(ops))

    def test_pdf_make_editable_not_in_catalog(self):
        ops = {item["operation"] for item in self.operations.list_operations("pdf")}
        self.assertNotIn("pdf_make_editable", ops)


class PdfToDocxTests(unittest.TestCase):
    def setUp(self):
        _install_handler_stubs()

        class FakeConverter:
            def __init__(self, input_path):
                self.input_path = input_path

            def convert(self, output_path, **kwargs):
                with open(output_path, "wb") as fh:
                    fh.write(_minimal_docx_bytes("Styled sample content"))

            def close(self):
                pass

        sys.modules["pdf2docx"] = types.SimpleNamespace(Converter=FakeConverter)
        sys.modules.pop("handlers.pdf_to_docx", None)
        self.mod = importlib.import_module("handlers.pdf_to_docx")

    def test_pdf_to_docx_uses_valid_docx_without_generated_label(self):
        result = self.mod._process(_minimal_pdf_bytes(), {"file_name": "sample.pdf"})

        self.assertTrue(zipfile.is_zipfile(io.BytesIO(result)))
        with zipfile.ZipFile(io.BytesIO(result)) as zf:
            content = "\n".join(zf.read(name).decode("utf-8", errors="ignore") for name in zf.namelist())
        self.assertIn("Styled sample content", content)
        self.assertNotIn("Converted Document", content)

    def test_pdf_to_docx_output_name_uses_input_basename(self):
        self.assertEqual(
            self.mod._output_filename({"file_name": "quarterly-report.pdf"}, "uploads/job/input.pdf"),
            "PDF to Word - quarterly-report.docx",
        )

    def test_pdf_to_docx_falls_back_to_text_when_all_converters_fail(self):
        sys.modules["pdf2docx"] = None
        self.mod._pdf_to_docx_libreoffice = lambda data: (_ for _ in ()).throw(RuntimeError("not found"))
        self.mod._extract_text_docx = lambda data: b"fallback-docx"
        try:
            self.assertEqual(self.mod._process(_minimal_pdf_bytes(), {}), b"fallback-docx")
        finally:
            sys.modules.pop("pdf2docx", None)

    def test_page_range_filters_pages(self):
        """Only the requested pages are extracted before conversion."""
        converted_page_counts = []

        class CountingConverter:
            def __init__(self, input_path):
                from pypdf import PdfReader
                self.n = len(PdfReader(input_path).pages)

            def convert(self, output_path, **kwargs):
                converted_page_counts.append(self.n)
                with open(output_path, "wb") as fh:
                    fh.write(_minimal_docx_bytes("pages"))

            def close(self):
                pass

        sys.modules["pdf2docx"] = types.SimpleNamespace(Converter=CountingConverter)
        pdf = _minimal_pdf_bytes(n_pages=10)
        self.mod._process(pdf, {"file_name": "doc.pdf", "params": {"page_range": "1-3"}})
        self.assertEqual(converted_page_counts, [3])

    def test_large_pdf_is_split_into_chunks_and_merged(self):
        """A PDF exceeding _CHUNK_SIZE pages is converted in chunks and merged."""
        chunk_page_counts = []

        def _real_docx_bytes(text):
            from docx import Document as DocxDocument
            doc = DocxDocument()
            doc.add_paragraph(text)
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()

        class CountingConverter:
            def __init__(self, input_path):
                from pypdf import PdfReader
                self.n = len(PdfReader(input_path).pages)

            def convert(self, output_path, **kwargs):
                chunk_page_counts.append(self.n)
                with open(output_path, "wb") as fh:
                    fh.write(_real_docx_bytes(f"chunk-{self.n}"))

            def close(self):
                pass

        sys.modules["pdf2docx"] = types.SimpleNamespace(Converter=CountingConverter)

        chunk_size = self.mod._CHUNK_SIZE
        total = chunk_size + 5
        pdf = _minimal_pdf_bytes(n_pages=total)
        result = self.mod._process(pdf, {"file_name": "big.pdf"})

        # Should have split into exactly 2 chunks
        self.assertEqual(len(chunk_page_counts), 2)
        self.assertEqual(chunk_page_counts[0], chunk_size)
        self.assertEqual(chunk_page_counts[1], 5)
        # Result must be a valid zip (DOCX)
        self.assertTrue(zipfile.is_zipfile(io.BytesIO(result)))

    def test_text_fallback_is_authoritative_and_never_switches_to_image_docx(self):
        pdf = _text_pdf_bytes()
        self.mod._pdf2docx_convert = lambda *_args, **_kwargs: _minimal_docx_bytes("")
        self.mod._build_text_fallback_docx = lambda *_args, **_kwargs: b"text-fallback"
        self.mod._build_image_fallback_docx = lambda *_args, **_kwargs: (_ for _ in ()).throw(
            AssertionError("image fallback must not be used when fallback_strategy=text")
        )

        out = self.mod._process(
            pdf,
            {"params": {"fallback_strategy": "text", "high_fidelity": False}},
        )
        self.assertEqual(out, b"text-fallback")

    def test_text_fallback_disables_hybrid_high_fidelity_override(self):
        pdf = _text_pdf_bytes()
        self.mod._hybrid_high_fidelity_docx = lambda *_args, **_kwargs: (_ for _ in ()).throw(
            AssertionError("hybrid mode must be disabled by strict text fallback")
        )
        self.mod._pdf2docx_convert = lambda *_args, **_kwargs: _minimal_docx_bytes("ok")

        out = self.mod._process(
            pdf,
            {"params": {"fallback_strategy": "text", "high_fidelity": True}},
        )
        self.assertTrue(zipfile.is_zipfile(io.BytesIO(out)))


class OfficeToPdfTests(unittest.TestCase):
    def setUp(self):
        self.dynamo, self.s3 = _install_handler_stubs()
        sys.modules.pop("handlers.docx_to_pdf", None)
        self.docx_mod = importlib.import_module("handlers.docx_to_pdf")

    def test_docx_to_pdf_handler_marks_failed_on_corrupt_document(self):
        calls = []
        self.dynamo.update_job = lambda *args, **kwargs: calls.append(("update", args, kwargs))
        self.dynamo.mark_failed = lambda *args, **kwargs: calls.append(("failed", args, kwargs))
        self.s3.get_bytes = lambda key: b"not a docx"
        self.docx_mod._docx_to_pdf = lambda data, paper_size=None: (_ for _ in ()).throw(ValueError("corrupt document"))

        event = {"Records": [{"body": json.dumps({"job_id": "job-1", "file_key": "uploads/job-1/bad.docx"})}]}
        with self.assertRaises(ValueError):
            self.docx_mod.handler(event, None)

        self.assertTrue(any(call[0] == "failed" and call[1][0] == "job-1" for call in calls))

    @unittest.skipUnless(shutil.which("libreoffice") or shutil.which("soffice"), "LibreOffice not installed")
    def test_docx_to_pdf_libreoffice_output_is_pdf(self):
        from docx import Document
        import io

        if shutil.which("soffice"):
            os.environ["LIBREOFFICE_BIN"] = shutil.which("soffice")
            self.docx_mod._LIBREOFFICE_BIN = os.environ["LIBREOFFICE_BIN"]

        doc = Document()
        doc.add_heading("Quarterly Report", level=1)
        doc.add_paragraph("Bold and table content")
        stream = io.BytesIO()
        doc.save(stream)

        result = self.docx_mod._docx_to_pdf(stream.getvalue())
        self.assertTrue(result.startswith(b"%PDF"))


class MarkdownConvertTests(unittest.TestCase):
    def setUp(self):
        self.dynamo, self.s3 = _install_handler_stubs()
        sys.modules.pop("handlers.markdown_convert", None)
        self.mod = importlib.import_module("handlers.markdown_convert")

    def test_invalid_target_format_fails_before_rendering(self):
        with self.assertRaises(ValueError) as ctx:
            self.mod.convert_markdown(b"# Notes", "rtf")
        self.assertIn("target_format must be one of", str(ctx.exception))

    def test_undecodable_input_fails_cleanly(self):
        with self.assertRaises(ValueError) as ctx:
            self.mod.convert_markdown(b"\xff\xfe\xff", "pdf")
        self.assertIn("valid UTF-8", str(ctx.exception))

    def test_output_filename_uses_input_basename(self):
        self.assertEqual(
            self.mod._output_filename({"file_name": "notes.md"}, "uploads/job/input.md", "pdf"),
            "Convert Markdown - notes.pdf",
        )
        self.assertEqual(
            self.mod._output_filename({"file_name": "notes.md"}, "uploads/job/input.md", "jpeg"),
            "Convert Markdown - notes.jpeg",
        )

    @unittest.skipUnless(
        importlib.util.find_spec("markdown")
        and importlib.util.find_spec("bs4")
        and importlib.util.find_spec("reportlab")
        and importlib.util.find_spec("docx")
        and importlib.util.find_spec("PIL"),
        "Markdown conversion dependencies not installed",
    )
    def test_markdown_outputs_are_valid_documents(self):
        from PIL import Image

        source = b"""# Title

Paragraph with **bold**, *italic*, [link](https://example.com), and `code`.

- one
- two

> quoted text

```python
print("ok")
```

| A | B |
| - | - |
| 1 | 2 |
"""

        pdf = self.mod.convert_markdown(source, "pdf")
        self.assertTrue(pdf.startswith(b"%PDF"))

        docx = self.mod.convert_markdown(source, "docx")
        self.assertTrue(zipfile.is_zipfile(__import__("io").BytesIO(docx)))
        with zipfile.ZipFile(__import__("io").BytesIO(docx)) as zf:
            self.assertIn("word/document.xml", zf.namelist())

        for target in ("png", "jpg", "tiff"):
            image_bytes = self.mod.convert_markdown(source, target)
            img = Image.open(__import__("io").BytesIO(image_bytes))
            self.assertGreater(img.width, 0)
            self.assertGreater(img.height, 0)

    @unittest.skipUnless(
        importlib.util.find_spec("markdown") and importlib.util.find_spec("bs4"),
        "Markdown parser dependencies not installed",
    )
    def test_parser_model_handles_common_markdown_blocks(self):
        blocks = self.mod.parse_markdown("""# Heading

Text with **bold** and *italic*.

1. First
2. Second

> Quote

```
code
```

| A | B |
| - | - |
| 1 | 2 |
""")
        kinds = [block.kind for block in blocks]
        self.assertIn("heading", kinds)
        self.assertIn("paragraph", kinds)
        self.assertIn("list", kinds)
        self.assertIn("quote", kinds)
        self.assertIn("code", kinds)
        self.assertIn("table", kinds)


if __name__ == "__main__":
    unittest.main()

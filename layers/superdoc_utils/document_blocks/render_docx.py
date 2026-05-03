from __future__ import annotations

import io

from .model import Block, Inline


def _add_docx_runs(paragraph, inlines: list[Inline]):
    for item in inlines:
        run = paragraph.add_run(item.text)
        run.bold = item.bold
        run.italic = item.italic
        if item.code:
            run.font.name = "Courier New"


def render_docx(blocks: list[Block]) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt

    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    def add_block(block: Block, *, quote: bool = False):
        if block.kind == "heading":
            paragraph = document.add_heading(level=min(max(block.level, 1), 3))
            _add_docx_runs(paragraph, block.inlines)
        elif block.kind == "paragraph":
            paragraph = document.add_paragraph()
            if quote:
                paragraph.paragraph_format.left_indent = Inches(0.25)
            _add_docx_runs(paragraph, block.inlines)
        elif block.kind == "list":
            style = "List Number" if block.ordered else "List Bullet"
            for item in block.items:
                _add_docx_runs(document.add_paragraph(style=style), item)
        elif block.kind == "quote":
            for nested in block.blocks:
                add_block(nested, quote=True)
        elif block.kind == "section":
            for nested in block.blocks:
                add_block(nested, quote=quote)
        elif block.kind == "code":
            paragraph = document.add_paragraph()
            if quote:
                paragraph.paragraph_format.left_indent = Inches(0.25)
            run = paragraph.add_run(block.text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif block.kind == "table":
            cols = max(len(block.headers), max((len(row) for row in block.rows), default=0))
            if cols == 0:
                return
            rows = len(block.rows) + (1 if block.headers else 0)
            table = document.add_table(rows=rows, cols=cols)
            table.style = "Table Grid"
            row_offset = 0
            if block.headers:
                for idx, cell in enumerate(block.headers):
                    _add_docx_runs(table.cell(0, idx).paragraphs[0], cell)
                    for run in table.cell(0, idx).paragraphs[0].runs:
                        run.bold = True
                row_offset = 1
            for row_idx, row in enumerate(block.rows):
                for col_idx, cell in enumerate(row):
                    _add_docx_runs(table.cell(row_idx + row_offset, col_idx).paragraphs[0], cell)

    for block in blocks:
        add_block(block)

    out = io.BytesIO()
    document.save(out)
    return out.getvalue()

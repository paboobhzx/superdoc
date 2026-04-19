#!/usr/bin/env python3
# =============================================================================
# SuperDoc Round 3a — End-to-End Smoke Test
# =============================================================================
#
# Covers:
#   - GET  /operations                        (catalog integrity + kind field)
#   - GET  /operations?input_type=<ext>       (filtering)
#   - POST /jobs                              (create_job for each backend_job op)
#   - PUT  <upload_url>                       (S3 presigned upload)
#   - GET  /jobs/<id>                         (polling → SUCCEEDED/FAILED)
#   - GET  <download_url>                     (result retrieval)
#   - GET  /files/download?key=<s3_key>       (3a-3 presign_download endpoint)
#   - POST /checkout                          (3a-2 dormant — expects 503)
#   - POST /stripe/webhook                    (3a-2 dormant — expects 503/400)
#
# Sample files generated in-memory:
#   - PDF   (ReportLab)
#   - DOCX  (python-docx)
#   - XLSX  (openpyxl)
#   - PNG   (Pillow)
#   - TXT   (stdlib)
#
# Does NOT cover (requires browser/DOM):
#   - Frontend pickerRouting.dispatchPick routing
#   - Editors auto-load via useKeyFileLoader hook
#   - OperationPicker UI + sessionStorage cache
#   - CloudFront cache/invalidation propagation
#
# Usage:
#   python3 round3a_smoke.py
#   python3 round3a_smoke.py --base-url https://... --verbose
#   python3 round3a_smoke.py --only pdf,docx       # subset by input type
#   python3 round3a_smoke.py --skip-paid-checks    # skip 3a-2 dormant checks
# =============================================================================

import argparse
import io
import json
import os
import re
import sys
import time
import traceback
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Tuple
from urllib import error as urlerror
from urllib import request as urlrequest

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

DEFAULT_BASE_URL = "https://g590x2ydn4.execute-api.us-east-1.amazonaws.com/dev"
POLL_INTERVAL_SEC = 2
POLL_TIMEOUT_SEC = 180  # 3 min per job
HTTP_TIMEOUT_SEC = 30

# ---------------------------------------------------------------------------
# Output helpers
# ---------------------------------------------------------------------------


class Colors:
    GREEN = "\033[92m"
    RED = "\033[91m"
    YELLOW = "\033[93m"
    BLUE = "\033[94m"
    MAGENTA = "\033[95m"
    CYAN = "\033[96m"
    GRAY = "\033[90m"
    BOLD = "\033[1m"
    RESET = "\033[0m"


def _supports_color() -> bool:
    if os.environ.get("NO_COLOR"):
        return False
    if not sys.stdout.isatty():
        return False
    return True


USE_COLOR = _supports_color()


def _c(color: str, text: str) -> str:
    if not USE_COLOR:
        return text
    return f"{color}{text}{Colors.RESET}"


def log_info(msg: str) -> None:
    print(f"{_c(Colors.BLUE, '[info]')} {msg}", flush=True)


def log_ok(msg: str) -> None:
    print(f"{_c(Colors.GREEN, '[ ok ]')} {msg}", flush=True)


def log_warn(msg: str) -> None:
    print(f"{_c(Colors.YELLOW, '[warn]')} {msg}", flush=True)


def log_fail(msg: str) -> None:
    print(f"{_c(Colors.RED, '[fail]')} {msg}", flush=True)


def log_step(msg: str) -> None:
    print(f"\n{_c(Colors.BOLD + Colors.CYAN, '▶ ' + msg)}", flush=True)


def log_verbose(msg: str, verbose: bool) -> None:
    if verbose:
        print(f"{_c(Colors.GRAY, '  · ' + msg)}", flush=True)


# ---------------------------------------------------------------------------
# Result tracking
# ---------------------------------------------------------------------------


@dataclass
class TestResult:
    name: str
    status: str  # "pass" | "fail" | "skip" | "warn"
    detail: str = ""
    duration_ms: int = 0


@dataclass
class Registry:
    results: List[TestResult] = field(default_factory=list)

    def record(self, name: str, status: str, detail: str = "", duration_ms: int = 0) -> None:
        self.results.append(TestResult(name, status, detail, duration_ms))
        emit = {
            "pass": log_ok,
            "fail": log_fail,
            "skip": log_warn,
            "warn": log_warn,
        }.get(status, log_info)
        suffix = ""
        if duration_ms:
            suffix = f" ({duration_ms}ms)"
        line = name
        if detail:
            line = f"{name} — {detail}"
        emit(f"{line}{suffix}")

    def summary(self) -> Tuple[int, int, int, int]:
        passed = sum(1 for r in self.results if r.status == "pass")
        failed = sum(1 for r in self.results if r.status == "fail")
        skipped = sum(1 for r in self.results if r.status == "skip")
        warned = sum(1 for r in self.results if r.status == "warn")
        return passed, failed, skipped, warned


# ---------------------------------------------------------------------------
# HTTP helpers (stdlib only — no requests dependency)
# ---------------------------------------------------------------------------


@dataclass
class HttpResponse:
    status: int
    headers: Dict[str, str]
    body: bytes

    def json(self) -> Any:
        if not self.body:
            return None
        return json.loads(self.body.decode("utf-8"))

    def text(self) -> str:
        return self.body.decode("utf-8", errors="replace")


def _build_multipart(fields: Dict[str, str], filename: str, content: bytes, content_type: str) -> Tuple[bytes, str]:
    """Build a multipart/form-data body for S3 presigned POST.
    The 'file' field MUST be last per S3 requirements."""
    import uuid as _uuid
    boundary = _uuid.uuid4().hex
    parts: List[bytes] = []
    for key, value in fields.items():
        parts.append(
            f"--{boundary}\r\n"
            f'Content-Disposition: form-data; name="{key}"\r\n\r\n'
            f"{value}\r\n".encode("utf-8")
        )
    # file field last
    parts.append(
        f"--{boundary}\r\n"
        f'Content-Disposition: form-data; name="file"; filename="{filename}"\r\n'
        f"Content-Type: {content_type}\r\n\r\n".encode("utf-8")
        + content
        + b"\r\n"
    )
    parts.append(f"--{boundary}--\r\n".encode("utf-8"))
    body = b"".join(parts)
    return body, f"multipart/form-data; boundary={boundary}"


def http_request(
    method: str,
    url: str,
    headers: Optional[Dict[str, str]] = None,
    body: Optional[bytes] = None,
    timeout: int = HTTP_TIMEOUT_SEC,
) -> HttpResponse:
    req = urlrequest.Request(url=url, method=method, data=body)
    if headers:
        for k, v in headers.items():
            req.add_header(k, v)
    try:
        with urlrequest.urlopen(req, timeout=timeout) as resp:
            raw = resp.read()
            hdrs = {k.lower(): v for k, v in resp.getheaders()}
            return HttpResponse(status=resp.status, headers=hdrs, body=raw)
    except urlerror.HTTPError as e:
        # Server responded with an error status — body may contain details.
        raw = b""
        try:
            raw = e.read()
        except Exception:
            pass
        hdrs = {}
        if e.headers:
            hdrs = {k.lower(): v for k, v in e.headers.items()}
        return HttpResponse(status=e.code, headers=hdrs, body=raw)
    except urlerror.URLError as e:
        # Network-level failure (DNS, connection refused, timeout, TLS).
        # Return synthetic status=0 so the test reports a failure instead of crashing.
        return HttpResponse(status=0, headers={}, body=f"network error: {e.reason}".encode("utf-8"))
    except (TimeoutError, ConnectionError, OSError) as e:
        return HttpResponse(status=0, headers={}, body=f"connection error: {e}".encode("utf-8"))


# ---------------------------------------------------------------------------
# Sample file generators (in-memory bytes)
# ---------------------------------------------------------------------------


def make_sample_txt() -> bytes:
    content = (
        "SuperDoc smoke test sample document.\n"
        "Line 2: Lorem ipsum dolor sit amet.\n"
        "Line 3: The quick brown fox jumps over the lazy dog.\n"
        "Line 4: 0123456789 !@#$%^&*()\n"
    )
    return content.encode("utf-8")


def make_sample_pdf() -> bytes:
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError:
        return _make_minimal_pdf_fallback()
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    c.setFont("Helvetica-Bold", 18)
    c.drawString(72, 720, "SuperDoc Smoke Test PDF")
    c.setFont("Helvetica", 12)
    c.drawString(72, 690, "Page 1 — lorem ipsum dolor sit amet.")
    c.drawString(72, 670, "Generated by round3a_smoke.py")
    c.showPage()
    c.setFont("Helvetica-Bold", 18)
    c.drawString(72, 720, "Page 2")
    c.setFont("Helvetica", 12)
    c.drawString(72, 690, "Second page for multi-page operations (split, pdf_to_image).")
    c.showPage()
    c.save()
    return buf.getvalue()


def _make_minimal_pdf_fallback() -> bytes:
    # Minimal valid PDF — single page, no text rendering. Enough to pass MIME check.
    pdf = (
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]/Contents 4 0 R>>endobj\n"
        b"4 0 obj<</Length 44>>stream\n"
        b"BT /F1 24 Tf 100 700 Td (SuperDoc) Tj ET\n"
        b"endstream\nendobj\n"
        b"xref\n0 5\n0000000000 65535 f\n"
        b"0000000009 00000 n\n"
        b"0000000052 00000 n\n"
        b"0000000098 00000 n\n"
        b"0000000165 00000 n\n"
        b"trailer<</Size 5/Root 1 0 R>>\n"
        b"startxref\n240\n%%EOF\n"
    )
    return pdf


def make_sample_docx() -> bytes:
    try:
        from docx import Document
    except ImportError:
        return _make_minimal_docx_fallback()
    doc = Document()
    doc.add_heading("SuperDoc Smoke Test", level=1)
    doc.add_paragraph("This is a sample DOCX generated for smoke testing.")
    doc.add_paragraph("Second paragraph — lorem ipsum dolor sit amet.")
    doc.add_heading("Section 2", level=2)
    doc.add_paragraph("Content for docx_to_txt extraction validation.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_minimal_docx_fallback() -> bytes:
    # Build a minimal valid DOCX (ZIP of OOXML) by hand — Python stdlib only.
    import zipfile

    content_types = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        "</Types>"
    )
    rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
        'Target="word/document.xml"/>'
        "</Relationships>"
    )
    document = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:body>"
        "<w:p><w:r><w:t>SuperDoc smoke test sample.</w:t></w:r></w:p>"
        "<w:p><w:r><w:t>Second paragraph for extraction validation.</w:t></w:r></w:p>"
        "</w:body>"
        "</w:document>"
    )
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", content_types)
        z.writestr("_rels/.rels", rels)
        z.writestr("word/document.xml", document)
    return buf.getvalue()


def make_sample_xlsx() -> bytes:
    try:
        from openpyxl import Workbook
    except ImportError:
        return _make_minimal_xlsx_fallback()
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["col_a", "col_b", "col_c"])
    ws.append([1, "hello", 3.14])
    ws.append([2, "world", 2.71])
    ws.append([3, "smoke", 1.41])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _make_minimal_xlsx_fallback() -> bytes:
    # Skip fallback for XLSX — the hand-built OOXML is error-prone; require openpyxl.
    raise RuntimeError("openpyxl required to generate XLSX sample — pip install openpyxl")


def make_sample_png() -> bytes:
    try:
        from PIL import Image, ImageDraw
    except ImportError:
        return _make_minimal_png_fallback()
    img = Image.new("RGB", (400, 300), color=(34, 139, 34))
    d = ImageDraw.Draw(img)
    d.text((20, 20), "SuperDoc", fill=(255, 255, 255))
    d.text((20, 60), "smoke test image", fill=(255, 255, 255))
    d.rectangle([(50, 100), (350, 250)], outline=(255, 255, 255), width=3)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _make_minimal_png_fallback() -> bytes:
    # 1x1 red PNG — 67 bytes, valid PNG.
    return bytes.fromhex(
        "89504e470d0a1a0a0000000d49484452000000010000000108060000001f15c489"
        "0000000d49444154789c6364f8cfc0000000030001018ad2a40f0000000049454e44ae426082"
    )


SAMPLE_GENERATORS = {
    "pdf": ("sample.pdf", "application/pdf", make_sample_pdf),
    "docx": (
        "sample.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        make_sample_docx,
    ),
    "xlsx": (
        "sample.xlsx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        make_sample_xlsx,
    ),
    "png": ("sample.png", "image/png", make_sample_png),
    "txt": ("sample.txt", "text/plain", make_sample_txt),
}


# ---------------------------------------------------------------------------
# SuperDoc API client
# ---------------------------------------------------------------------------


class SuperDocClient:
    def __init__(self, base_url: str, verbose: bool = False):
        self.base_url = base_url.rstrip("/")
        self.verbose = verbose

    def _url(self, path: str) -> str:
        if not path.startswith("/"):
            path = "/" + path
        return f"{self.base_url}{path}"

    def get_operations(self, input_type: Optional[str] = None) -> HttpResponse:
        path = "/operations"
        if input_type:
            path = f"/operations?input_type={input_type}"
        log_verbose(f"GET {path}", self.verbose)
        return http_request("GET", self._url(path))

    def create_job(
        self,
        operation: str,
        filename: str,
        file_size_bytes: int = 0,
        session_id: Optional[str] = None,
        params: Optional[Dict[str, Any]] = None,
    ) -> HttpResponse:
        payload: Dict[str, Any] = {
            "operation": operation,
            "filename": filename,
            "file_size_bytes": file_size_bytes,
        }
        if session_id:
            payload["session_id"] = session_id
        if params:
            payload["params"] = params
        body = json.dumps(payload).encode("utf-8")
        log_verbose(f"POST /jobs body={payload}", self.verbose)
        return http_request(
            "POST",
            self._url("/jobs"),
            headers={"Content-Type": "application/json"},
            body=body,
        )

    def get_job(self, job_id: str) -> HttpResponse:
        log_verbose(f"GET /jobs/{job_id}", self.verbose)
        return http_request("GET", self._url(f"/jobs/{job_id}"))

    def upload_s3(self, upload_url: str, content: bytes, content_type: str) -> HttpResponse:
        log_verbose(f"PUT {upload_url[:80]}... ({len(content)} bytes)", self.verbose)
        return http_request(
            "PUT",
            upload_url,
            headers={"Content-Type": content_type},
            body=content,
            timeout=60,
        )

    def download(self, url: str) -> HttpResponse:
        log_verbose(f"GET {url[:80]}...", self.verbose)
        return http_request("GET", url, timeout=60)

    def presign_download(self, s3_key: str) -> HttpResponse:
        from urllib.parse import quote

        path = f"/files/download?key={quote(s3_key, safe='/')}"
        log_verbose(f"GET {path}", self.verbose)
        return http_request("GET", self._url(path))

    def create_checkout(self, payload: Dict[str, Any]) -> HttpResponse:
        body = json.dumps(payload).encode("utf-8")
        log_verbose(f"POST /checkout body={payload}", self.verbose)
        return http_request(
            "POST",
            self._url("/checkout"),
            headers={"Content-Type": "application/json"},
            body=body,
        )

    def process_job(self, job_id: str) -> HttpResponse:
        log_verbose(f"POST /jobs/{job_id}/process", self.verbose)
        return http_request(
            "POST",
            self._url(f"/jobs/{job_id}/process"),
            headers={"Content-Type": "application/json"},
            body=b"{}",
        )

    def post_stripe_webhook(self, body: bytes, signature: str) -> HttpResponse:
        log_verbose(f"POST /stripe/webhook sig={signature[:30]}...", self.verbose)
        return http_request(
            "POST",
            self._url("/stripe/webhook"),
            headers={
                "Content-Type": "application/json",
                "Stripe-Signature": signature,
            },
            body=body,
        )


# ---------------------------------------------------------------------------
# Test: catalog integrity
# ---------------------------------------------------------------------------


def test_catalog_integrity(client: SuperDocClient, registry: Registry) -> Optional[List[Dict[str, Any]]]:
    log_step("Catalog integrity — GET /operations")
    t0 = time.time()
    resp = client.get_operations()
    dur = int((time.time() - t0) * 1000)

    if resp.status != 200:
        registry.record(
            "GET /operations returns 200",
            "fail",
            f"got {resp.status}: {resp.text()[:200]}",
            dur,
        )
        return None

    try:
        data = resp.json()
    except Exception as e:
        registry.record("GET /operations returns JSON", "fail", str(e), dur)
        return None

    # Some responses wrap the list in {"operations": [...]}, others return a bare list.
    if isinstance(data, dict):
        ops = data.get("operations")
    else:
        ops = data
    if not isinstance(ops, list):
        registry.record(
            "GET /operations shape",
            "fail",
            f"expected list or {{operations: [...]}}; got {type(data).__name__}",
            dur,
        )
        return None

    registry.record(
        "GET /operations returns 200 with list",
        "pass",
        f"{len(ops)} ops",
        dur,
    )

    # kind field presence
    missing_kind = [o.get("id", "?") for o in ops if "kind" not in o]
    if missing_kind:
        registry.record(
            "All ops have 'kind' field",
            "fail",
            f"missing in: {missing_kind}",
        )
    else:
        registry.record("All ops have 'kind' field", "pass", f"{len(ops)} ops checked")

    # kind values valid
    valid_kinds = {"backend_job", "client_editor", "paid_backend_job"}
    bad_kinds = [(o.get("operation"), o.get("kind")) for o in ops if o.get("kind") not in valid_kinds]
    if bad_kinds:
        registry.record("All 'kind' values are valid", "fail", f"bad: {bad_kinds}")
    else:
        registry.record("All 'kind' values are valid", "pass", f"values in {valid_kinds}")

    # 4 new ops from 3a-1 present
    new_ops = {"pdf_to_image", "image_to_pdf", "docx_to_txt", "xlsx_to_csv"}
    present_ids = {o.get("operation") for o in ops}
    missing_new = new_ops - present_ids
    if missing_new:
        registry.record("Round 3a-1 new ops present", "fail", f"missing: {missing_new}")
    else:
        registry.record("Round 3a-1 new ops present", "pass", f"{sorted(new_ops)}")

    # dormant check: no paid_backend_job yet (3a-2 is dormant)
    paid_ops = [o.get("operation") for o in ops if o.get("kind") == "paid_backend_job"]
    if paid_ops:
        registry.record(
            "Stripe dormant state (no paid ops yet)",
            "warn",
            f"found paid ops: {paid_ops} — either Round 3b already shipped or misconfiguration",
        )
    else:
        registry.record("Stripe dormant state (no paid ops yet)", "pass", "0 paid_backend_job ops")

    # doc_edit is client_editor (critical for 3a-3 routing)
    doc_edit = next((o for o in ops if o.get("operation") == "doc_edit"), None)
    if doc_edit is None:
        registry.record("doc_edit op exists", "warn", "not found in catalog")
    elif doc_edit.get("kind") != "client_editor":
        registry.record(
            "doc_edit has kind=client_editor",
            "fail",
            f"got kind={doc_edit.get('kind')}",
        )
    else:
        registry.record("doc_edit has kind=client_editor", "pass")

    return ops


# ---------------------------------------------------------------------------
# Test: catalog filters by input_type
# ---------------------------------------------------------------------------


def test_catalog_filters(client: SuperDocClient, registry: Registry, all_ops: List[Dict[str, Any]]) -> None:
    log_step("Catalog filters — GET /operations?input_type=<ext>")
    # Expected containment rules: new ops appear in their input_type filter
    expectations = {
        "pdf": "pdf_to_image",
        "docx": "docx_to_txt",
        "xlsx": "xlsx_to_csv",
        "png": "image_to_pdf",
    }
    for input_type, expected_id in expectations.items():
        t0 = time.time()
        resp = client.get_operations(input_type=input_type)
        dur = int((time.time() - t0) * 1000)
        if resp.status != 200:
            registry.record(
                f"GET /operations?input_type={input_type}",
                "fail",
                f"status {resp.status}",
                dur,
            )
            continue
        try:
            data = resp.json()
        except Exception as e:
            registry.record(
                f"GET /operations?input_type={input_type} JSON",
                "fail",
                str(e),
                dur,
            )
            continue
        if isinstance(data, dict):
            ops = data.get("operations")
        else:
            ops = data
        if not isinstance(ops, list):
            registry.record(
                f"GET /operations?input_type={input_type} shape",
                "fail",
                "not a list",
                dur,
            )
            continue
        ids = {o.get("operation") for o in ops}
        if expected_id in ids:
            registry.record(
                f"input_type={input_type} contains {expected_id}",
                "pass",
                f"{len(ops)} ops",
                dur,
            )
        else:
            registry.record(
                f"input_type={input_type} contains {expected_id}",
                "fail",
                f"got {sorted(ids)}",
                dur,
            )


# ---------------------------------------------------------------------------
# Test: full backend_job flow (one op per input_type)
# ---------------------------------------------------------------------------


def poll_job(client: SuperDocClient, job_id: str, timeout_sec: int = POLL_TIMEOUT_SEC) -> Dict[str, Any]:
    deadline = time.time() + timeout_sec
    last_status = None
    while time.time() < deadline:
        resp = client.get_job(job_id)
        if resp.status != 200:
            time.sleep(POLL_INTERVAL_SEC)
            continue
        data = resp.json()
        status = data.get("status")
        if status != last_status:
            log_verbose(f"job {job_id} status={status}", client.verbose)
            last_status = status
        if status in ("DONE", "SUCCEEDED", "FAILED", "ERROR"):
            return data
        time.sleep(POLL_INTERVAL_SEC)
    raise TimeoutError(f"job {job_id} did not finish within {timeout_sec}s (last status={last_status})")


def run_backend_job(
    client: SuperDocClient,
    registry: Registry,
    operation: str,
    input_type: str,
    params: Optional[Dict[str, Any]] = None,
    session_id: Optional[str] = None,
) -> Optional[str]:
    """Returns S3 key of uploaded input on success, for reuse in presign test."""
    test_name = f"backend_job: {operation} ({input_type})"
    log_step(test_name)
    if input_type not in SAMPLE_GENERATORS:
        registry.record(test_name, "skip", f"no generator for input_type={input_type}")
        return None

    filename, content_type, generator = SAMPLE_GENERATORS[input_type]
    try:
        content = generator()
    except Exception as e:
        registry.record(test_name, "fail", f"sample generation failed: {e}")
        return None

    t0 = time.time()

    # 1. Create job
    resp = client.create_job(operation, filename, file_size_bytes=len(content), session_id=session_id, params=params)
    if resp.status not in (200, 201):
        registry.record(
            f"{test_name} :: create_job",
            "fail",
            f"status {resp.status}: {resp.text()[:200]}",
        )
        return None
    job = resp.json()
    job_id = job.get("job_id") or job.get("id")
    s3_key = job.get("file_key") or job.get("s3_key") or job.get("key") or job.get("input_key")

    # API returns presigned POST: {"upload": {"url": "...", "fields": {...}}, ...}
    upload_obj = job.get("upload") or {}
    upload_url = upload_obj.get("url") or job.get("upload_url") or job.get("uploadUrl")
    upload_fields: Dict[str, str] = upload_obj.get("fields") or {}

    if not job_id or not upload_url:
        registry.record(
            f"{test_name} :: create_job response shape",
            "fail",
            f"missing job_id or upload_url: {job}",
        )
        return None
    log_verbose(f"job_id={job_id} s3_key={s3_key}", client.verbose)

    # 2. Upload input — presigned POST (multipart) or presigned PUT
    if upload_fields:
        body, ct = _build_multipart(upload_fields, filename, content, content_type)
        resp = http_request("POST", upload_url, headers={"Content-Type": ct}, body=body, timeout=60)
        ok_statuses = (200, 201, 204)
    else:
        resp = client.upload_s3(upload_url, content, content_type)
        ok_statuses = (200, 204)
    if resp.status not in ok_statuses:
        registry.record(
            f"{test_name} :: upload",
            "fail",
            f"S3 upload status {resp.status}: {resp.text()[:200]}",
        )
        return None

    # 3. Trigger processing — POST /jobs/{jobId}/process → QUEUED → SQS → worker
    resp = client.process_job(job_id)
    if resp.status not in (200, 202):
        registry.record(
            f"{test_name} :: process_job",
            "fail",
            f"status {resp.status}: {resp.text()[:200]}",
        )
        return s3_key

    # 4. Poll
    try:
        final = poll_job(client, job_id)
    except TimeoutError as e:
        registry.record(f"{test_name} :: poll", "fail", str(e))
        return s3_key

    status = final.get("status")
    dur = int((time.time() - t0) * 1000)
    if status not in ("DONE", "SUCCEEDED"):
        err_detail = final.get("error") or final.get("error_message") or final.get("detail") or ""
        registry.record(
            test_name,
            "fail",
            f"final status={status} error={err_detail}",
            dur,
        )
        return s3_key

    # 4. Verify result download URL works
    result_url = final.get("download_url") or final.get("result_url") or final.get("output_url")
    if not result_url:
        registry.record(
            test_name,
            "warn",
            "job succeeded but no download_url in response",
            dur,
        )
        return s3_key

    resp = client.download(result_url)
    if resp.status != 200 or len(resp.body) == 0:
        registry.record(
            test_name,
            "fail",
            f"download status {resp.status} size {len(resp.body)}",
            dur,
        )
        return s3_key

    registry.record(
        test_name,
        "pass",
        f"output {len(resp.body)} bytes",
        dur,
    )
    return s3_key


# ---------------------------------------------------------------------------
# Test: /files/download presign endpoint (3a-3)
# ---------------------------------------------------------------------------


def test_presign_download(client: SuperDocClient, registry: Registry, known_key: Optional[str]) -> None:
    log_step("Presign download endpoint — GET /files/download (3a-3)")

    # Case 1: missing ?key= param — expect 400
    t0 = time.time()
    from urllib.parse import quote  # noqa

    resp = http_request("GET", client._url("/files/download"))
    dur = int((time.time() - t0) * 1000)
    if resp.status == 400:
        registry.record("presign: no key → 400", "pass", "", dur)
    elif resp.status == 200:
        registry.record("presign: no key → 400", "fail", "got 200 (should reject)", dur)
    else:
        registry.record(
            "presign: no key → 400",
            "warn",
            f"got {resp.status} (expected 400)",
            dur,
        )

    # Case 2: traversal attempt — expect 400/403
    t0 = time.time()
    resp = client.presign_download("../../etc/passwd")
    dur = int((time.time() - t0) * 1000)
    if resp.status in (400, 403):
        registry.record(f"presign: traversal → {resp.status}", "pass", "", dur)
    else:
        registry.record(
            "presign: traversal rejected",
            "fail",
            f"got {resp.status} (expected 400/403)",
            dur,
        )

    # Case 3: key outside whitelist (uploads/* or users/*) — expect 400/403
    t0 = time.time()
    resp = client.presign_download("secrets/admin.txt")
    dur = int((time.time() - t0) * 1000)
    if resp.status in (400, 403):
        registry.record(
            f"presign: non-whitelisted prefix → {resp.status}",
            "pass",
            "",
            dur,
        )
    else:
        registry.record(
            "presign: non-whitelisted prefix rejected",
            "warn",
            f"got {resp.status} (expected 400/403)",
            dur,
        )

    # Case 4: valid whitelisted key — expect 200 with presigned URL
    if not known_key:
        registry.record(
            "presign: valid key → 200 with URL",
            "skip",
            "no known key from prior uploads (run backend_job tests first)",
        )
        return

    if not (known_key.startswith("uploads/") or known_key.startswith("users/")):
        registry.record(
            "presign: valid key → 200 with URL",
            "skip",
            f"known key {known_key} not in whitelisted prefix",
        )
        return

    t0 = time.time()
    resp = client.presign_download(known_key)
    dur = int((time.time() - t0) * 1000)
    if resp.status != 200:
        registry.record(
            "presign: valid key → 200",
            "fail",
            f"status {resp.status}: {resp.text()[:200]}",
            dur,
        )
        return

    try:
        data = resp.json()
    except Exception as e:
        registry.record("presign: valid key → JSON", "fail", str(e), dur)
        return

    url = data.get("url") or data.get("download_url") or data.get("presigned_url")
    if not url:
        registry.record(
            "presign: response has URL",
            "fail",
            f"no url field in: {list(data.keys())}",
            dur,
        )
        return

    # URL should be a legitimate S3 presigned URL
    if "X-Amz-Signature" not in url and "X-Amz-Algorithm" not in url:
        registry.record(
            "presign: URL is S3 presigned",
            "warn",
            "no X-Amz-Signature in URL",
            dur,
        )
    else:
        registry.record("presign: URL is S3 presigned", "pass", "", dur)

    # Fetch it and verify bytes come back
    resp2 = client.download(url)
    if resp2.status == 200 and len(resp2.body) > 0:
        registry.record(
            "presign: presigned URL downloads successfully",
            "pass",
            f"{len(resp2.body)} bytes",
        )
    else:
        registry.record(
            "presign: presigned URL downloads successfully",
            "fail",
            f"status {resp2.status} size {len(resp2.body)}",
        )


# ---------------------------------------------------------------------------
# Test: Stripe dormant checks (3a-2)
# ---------------------------------------------------------------------------


def test_stripe_dormant(client: SuperDocClient, registry: Registry) -> None:
    log_step("Stripe dormant infra — 3a-2")

    # /checkout with placeholder SSM values — expect 503 "Payments not configured"
    t0 = time.time()
    resp = client.create_checkout({
        "operation": "pdf_to_docx",
        "file_name": "test.pdf",
        "file_size_bytes": 1024,
        "success_url": "https://superdoc.pablobhz.cloud/success",
        "cancel_url": "https://superdoc.pablobhz.cloud/cancel",
    })
    dur = int((time.time() - t0) * 1000)
    if resp.status == 503:
        registry.record(
            "POST /checkout → 503 (dormant, placeholder SSM)",
            "pass",
            "",
            dur,
        )
    elif resp.status == 200:
        registry.record(
            "POST /checkout → 503 (dormant, placeholder SSM)",
            "warn",
            "got 200 — Stripe may already be configured (real SSM secrets set)",
            dur,
        )
    else:
        registry.record(
            "POST /checkout → 503 (dormant, placeholder SSM)",
            "warn",
            f"got {resp.status}: {resp.text()[:200]}",
            dur,
        )

    # /stripe/webhook with bogus signature — expect 400/401/503 (reject)
    t0 = time.time()
    body = b'{"type":"checkout.session.completed","data":{"object":{}}}'
    resp = client.post_stripe_webhook(body, "t=0,v1=bogus")
    dur = int((time.time() - t0) * 1000)
    if resp.status in (400, 401, 403, 503):
        registry.record(
            f"POST /stripe/webhook bogus sig → {resp.status} (reject)",
            "pass",
            "",
            dur,
        )
    elif resp.status == 200:
        registry.record(
            "POST /stripe/webhook bogus sig → reject",
            "fail",
            "got 200 — webhook accepts unsigned requests (security hole)",
            dur,
        )
    else:
        registry.record(
            "POST /stripe/webhook bogus sig → reject",
            "warn",
            f"got {resp.status}",
            dur,
        )


# ---------------------------------------------------------------------------
# Orchestration
# ---------------------------------------------------------------------------


def pick_backend_ops_for_smoke(ops: List[Dict[str, Any]], only: Optional[set]) -> List[Tuple[str, str, Optional[Dict[str, Any]]]]:
    """Pick one representative backend_job op per input_type. Prefer the new 3a-1 ops."""
    priority = ["pdf_to_image", "docx_to_txt", "xlsx_to_csv", "image_to_pdf"]
    backend_ops = [o for o in ops if o.get("kind") == "backend_job"]

    picked: Dict[str, Tuple[str, str, Optional[Dict[str, Any]]]] = {}
    # First pass: priority ops
    for op in backend_ops:
        op_id = op.get("operation")
        if op_id not in priority:
            continue
        input_types = op.get("input_types") or []
        if not input_types:
            continue
        input_type = input_types[0]
        if only and input_type not in only:
            continue
        if input_type not in picked:
            picked[input_type] = (op_id, input_type, None)
    # Second pass: fill remaining input_types with any backend_job op
    for op in backend_ops:
        input_types = op.get("input_types") or []
        for it in input_types:
            if only and it not in only:
                continue
            if it not in picked and it in SAMPLE_GENERATORS:
                picked[it] = (op.get("operation"), it, None)
                break
    return list(picked.values())


def main() -> int:
    parser = argparse.ArgumentParser(
        description="SuperDoc Round 3a automated smoke test",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("--base-url", default=DEFAULT_BASE_URL, help=f"API base URL (default: {DEFAULT_BASE_URL})")
    parser.add_argument("--verbose", "-v", action="store_true", help="Verbose logging")
    parser.add_argument(
        "--only",
        default="",
        help="Comma-separated input types to test (pdf,docx,xlsx,png). Default: all",
    )
    parser.add_argument("--skip-paid-checks", action="store_true", help="Skip 3a-2 dormant checks")
    parser.add_argument("--skip-backend-jobs", action="store_true", help="Skip full backend_job flows (catalog + presign only)")
    args = parser.parse_args()

    only_set = None
    if args.only:
        only_set = {s.strip() for s in args.only.split(",") if s.strip()}

    print(_c(Colors.BOLD + Colors.MAGENTA, "\n╔════════════════════════════════════════════════════════════╗"))
    print(_c(Colors.BOLD + Colors.MAGENTA, "║         SuperDoc Round 3a — End-to-End Smoke Test          ║"))
    print(_c(Colors.BOLD + Colors.MAGENTA, "╚════════════════════════════════════════════════════════════╝"))
    print(f"Base URL:  {args.base_url}")
    print(f"Verbose:   {args.verbose}")
    print(f"Only:      {only_set or 'all'}")
    print(f"Skip paid: {args.skip_paid_checks}")
    print()

    import uuid as _uuid
    smoke_session = f"smoke-{_uuid.uuid4().hex[:12]}"
    log_info(f"Smoke session_id: {smoke_session}")

    client = SuperDocClient(args.base_url, verbose=args.verbose)
    registry = Registry()

    # 1. Catalog integrity
    ops = test_catalog_integrity(client, registry)
    if ops is None:
        print_summary(registry)
        return 1

    # 2. Catalog filters
    test_catalog_filters(client, registry, ops)

    # 3. Backend job flows — one per input_type
    known_key: Optional[str] = None
    if not args.skip_backend_jobs:
        picks = pick_backend_ops_for_smoke(ops, only_set)
        if not picks:
            registry.record(
                "backend_job flow selection",
                "warn",
                "no backend_job ops matched filter",
            )
        for op_id, input_type, params in picks:
            key = run_backend_job(client, registry, op_id, input_type, params=params, session_id=smoke_session)
            if key and not known_key:
                known_key = key
    else:
        registry.record("backend_job flows", "skip", "--skip-backend-jobs set")

    # 4. Presign download (3a-3)
    test_presign_download(client, registry, known_key)

    # 5. Stripe dormant (3a-2)
    if args.skip_paid_checks:
        registry.record("Stripe dormant checks", "skip", "--skip-paid-checks set")
    else:
        test_stripe_dormant(client, registry)

    return print_summary(registry)


def print_summary(registry: Registry) -> int:
    print()
    print(_c(Colors.BOLD, "═" * 62))
    print(_c(Colors.BOLD, " RESULT MATRIX"))
    print(_c(Colors.BOLD, "═" * 62))
    for r in registry.results:
        icon = {
            "pass": _c(Colors.GREEN, "✓"),
            "fail": _c(Colors.RED, "✗"),
            "skip": _c(Colors.YELLOW, "—"),
            "warn": _c(Colors.YELLOW, "!"),
        }.get(r.status, "?")
        suffix = ""
        if r.duration_ms:
            suffix = f"  {_c(Colors.GRAY, f'{r.duration_ms}ms')}"
        line = f"  {icon}  {r.name}"
        if r.detail:
            line += f"  {_c(Colors.GRAY, '— ' + r.detail)}"
        line += suffix
        print(line)

    passed, failed, skipped, warned = registry.summary()
    print()
    print(_c(Colors.BOLD, "═" * 62))

    # Colorize counts only when nonzero so 0-fail line doesn't look alarming.
    passed_str = _c(Colors.GREEN, f"{passed} passed")
    if failed:
        failed_str = _c(Colors.RED, f"{failed} failed")
    else:
        failed_str = f"{failed} failed"
    if warned:
        warned_str = _c(Colors.YELLOW, f"{warned} warn")
    else:
        warned_str = f"{warned} warn"
    if skipped:
        skipped_str = _c(Colors.YELLOW, f"{skipped} skipped")
    else:
        skipped_str = f"{skipped} skipped"
    summary_parts = [passed_str, failed_str, warned_str, skipped_str]
    print(" " + "  ·  ".join(summary_parts))
    print(_c(Colors.BOLD, "═" * 62))
    print()

    if failed > 0:
        print(_c(Colors.RED + Colors.BOLD, "SMOKE TEST FAILED"))
        print()
        print(_c(Colors.BOLD, "Manual validation still required (requires browser DOM):"))
        print("  · Frontend OperationPicker UI rendering")
        print("  · dispatchPick routing by kind (client_editor → /<type>-editor?key=)")
        print("  · useKeyFileLoader auto-loading file in each editor")
        print("  · sessionStorage cache on /operations")
        return 1

    print(_c(Colors.GREEN + Colors.BOLD, "SMOKE TEST PASSED"))
    print()
    print(_c(Colors.BOLD, "Manual validation still required (requires browser DOM):"))
    print("  · Open https://superdoc.pablobhz.cloud in browser (hard refresh)")
    print("  · Drop a .docx → pick 'doc_edit' → verify redirect to /docx-editor?key=...")
    print("  · Verify editor auto-loads file (no manual file picker)")
    print("  · Repeat for .xlsx, .png")
    print("  · DevTools Network: confirm /files/download called + S3 blob fetched")
    return 0


if __name__ == "__main__":
    try:
        sys.exit(main())
    except KeyboardInterrupt:
        print("\nInterrupted.", file=sys.stderr)
        sys.exit(130)
    except Exception:
        traceback.print_exc()
        sys.exit(2)

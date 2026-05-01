from __future__ import annotations

import html
import io
import json
import os
from dataclasses import dataclass, field
from typing import Iterable

import dynamo
import s3
from logger import get_logger

log = get_logger(__name__)

TARGET_FORMATS = {"pdf", "docx", "png", "jpg", "jpeg", "tiff"}
IMAGE_FORMATS = {"png", "jpg", "jpeg", "tiff"}
IMAGE_WIDTH = int(os.environ.get("MARKDOWN_IMAGE_WIDTH", "1200"))
IMAGE_MAX_HEIGHT = int(os.environ.get("MARKDOWN_IMAGE_MAX_HEIGHT", "65000"))
IMAGE_MAX_PIXELS = int(os.environ.get("MARKDOWN_IMAGE_MAX_PIXELS", "50000000"))


@dataclass
class Inline:
    text: str
    bold: bool = False
    italic: bool = False
    code: bool = False
    href: str | None = None


@dataclass
class Block:
    kind: str
    inlines: list[Inline] = field(default_factory=list)
    level: int = 0
    ordered: bool = False
    items: list[list[Inline]] = field(default_factory=list)
    text: str = ""
    headers: list[list[Inline]] = field(default_factory=list)
    rows: list[list[list[Inline]]] = field(default_factory=list)
    blocks: list["Block"] = field(default_factory=list)


def _decode_text(data: bytes) -> str:
    try:
        return data.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise ValueError("Input must be valid UTF-8 Markdown or plain text") from exc


def _node_text(inlines: Iterable[Inline]) -> str:
    return "".join(part.text for part in inlines)


def _parse_inlines(node, *, bold: bool = False, italic: bool = False, code: bool = False, href: str | None = None) -> list[Inline]:
    from bs4 import NavigableString

    if isinstance(node, NavigableString):
        text = str(node)
        return [Inline(text=text, bold=bold, italic=italic, code=code, href=href)] if text else []

    name = getattr(node, "name", "")
    next_bold = bold or name in ("strong", "b")
    next_italic = italic or name in ("em", "i")
    next_code = code or name == "code"
    next_href = href
    if name == "a":
        next_href = node.get("href") or href
    if name == "br":
        return [Inline("\n", bold=bold, italic=italic, code=code, href=href)]

    parts: list[Inline] = []
    for child in getattr(node, "children", []):
        parts.extend(_parse_inlines(child, bold=next_bold, italic=next_italic, code=next_code, href=next_href))
    return parts


def parse_markdown(text: str) -> list[Block]:
    import markdown
    from bs4 import BeautifulSoup, NavigableString

    md = markdown.Markdown(extensions=["extra", "sane_lists"])
    soup = BeautifulSoup(md.convert(text), "html.parser")
    blocks: list[Block] = []

    def parse_block(node) -> Block | None:
        if isinstance(node, NavigableString):
            if str(node).strip():
                return Block(kind="paragraph", inlines=[Inline(str(node).strip())])
            return None

        name = getattr(node, "name", "")
        if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            return Block(kind="heading", level=int(name[1]), inlines=_parse_inlines(node))
        if name == "p":
            return Block(kind="paragraph", inlines=_parse_inlines(node))
        if name in ("ul", "ol"):
            items: list[list[Inline]] = []
            for li in node.find_all("li", recursive=False):
                items.append(_parse_inlines(li))
            return Block(kind="list", ordered=name == "ol", items=items)
        if name == "blockquote":
            quote_blocks = [parsed for child in node.children if (parsed := parse_block(child)) is not None]
            return Block(kind="quote", blocks=quote_blocks)
        if name == "pre":
            return Block(kind="code", text=node.get_text().rstrip("\n"))
        if name == "table":
            headers: list[list[Inline]] = []
            rows: list[list[list[Inline]]] = []
            thead = node.find("thead")
            if thead:
                first = thead.find("tr")
                if first:
                    headers = [_parse_inlines(cell) for cell in first.find_all(["th", "td"], recursive=False)]
            tbody = node.find("tbody") or node
            for tr in tbody.find_all("tr", recursive=False):
                cells = [_parse_inlines(cell) for cell in tr.find_all(["td", "th"], recursive=False)]
                if cells:
                    rows.append(cells)
            if not headers and rows:
                headers = rows.pop(0)
            return Block(kind="table", headers=headers, rows=rows)
        return None

    for child in soup.children:
        block = parse_block(child)
        if block is not None:
            blocks.append(block)
    return blocks or [Block(kind="paragraph", inlines=[Inline("")])]


def _paragraph_markup(inlines: list[Inline]) -> str:
    parts: list[str] = []
    for item in inlines:
        text = html.escape(item.text).replace("\n", "<br/>")
        if not text:
            continue
        if item.code:
            text = f'<font name="Courier">{text}</font>'
        if item.bold:
            text = f"<b>{text}</b>"
        if item.italic:
            text = f"<i>{text}</i>"
        if item.href:
            link = html.escape(item.href, quote=True)
            text = f'<link href="{link}" color="blue">{text}</link>'
        parts.append(text)
    return "".join(parts) or "&nbsp;"


def render_pdf(blocks: list[Block]) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import ListFlowable, ListItem, Paragraph, Preformatted, SimpleDocTemplate, Spacer, Table, TableStyle

    styles = getSampleStyleSheet()
    styles["Normal"].fontName = "Helvetica"
    styles["Normal"].fontSize = 10.5
    styles["Normal"].leading = 15
    styles["Code"].fontName = "Courier"
    styles["Code"].fontSize = 9
    styles["Code"].leading = 12

    story = []

    def add_block(block: Block, *, quote: bool = False):
        left_indent = 18 if quote else 0
        if block.kind == "heading":
            style = styles[f"Heading{min(max(block.level, 1), 3)}"]
            story.append(Paragraph(_paragraph_markup(block.inlines), style))
            story.append(Spacer(1, 6))
        elif block.kind == "paragraph":
            style = styles["Normal"].clone("QuoteParagraph" if quote else "BodyParagraph")
            style.leftIndent = left_indent
            story.append(Paragraph(_paragraph_markup(block.inlines), style))
            story.append(Spacer(1, 8))
        elif block.kind == "list":
            items = [ListItem(Paragraph(_paragraph_markup(item), styles["Normal"])) for item in block.items]
            story.append(ListFlowable(items, bulletType="1" if block.ordered else "bullet", leftIndent=24 + left_indent))
            story.append(Spacer(1, 8))
        elif block.kind == "quote":
            for nested in block.blocks:
                add_block(nested, quote=True)
        elif block.kind == "code":
            story.append(Preformatted(block.text or " ", styles["Code"], maxLineLength=90))
            story.append(Spacer(1, 8))
        elif block.kind == "table":
            data = []
            if block.headers:
                data.append([Paragraph(_paragraph_markup(cell), styles["Normal"]) for cell in block.headers])
            for row in block.rows:
                data.append([Paragraph(_paragraph_markup(cell), styles["Normal"]) for cell in row])
            if data:
                table = Table(data, hAlign="LEFT")
                table.setStyle(TableStyle([
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D0D7DE")),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F6F8FA")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]))
                story.append(table)
                story.append(Spacer(1, 10))

    for block in blocks:
        add_block(block)

    out = io.BytesIO()
    def paint_background(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(colors.white)
        canvas.rect(0, 0, doc.pagesize[0], doc.pagesize[1], stroke=0, fill=1)
        canvas.restoreState()

    doc = SimpleDocTemplate(out, pagesize=letter, leftMargin=0.8 * inch, rightMargin=0.8 * inch, topMargin=0.7 * inch, bottomMargin=0.7 * inch)
    doc.build(story, onFirstPage=paint_background, onLaterPages=paint_background)
    return out.getvalue()


def _add_docx_runs(paragraph, inlines: list[Inline]):
    for item in inlines:
        run = paragraph.add_run(item.text)
        run.bold = item.bold
        run.italic = item.italic
        if item.code:
            run.font.name = "Courier New"


def render_docx(blocks: list[Block]) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt

    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    for block in blocks:
        if block.kind == "heading":
            paragraph = document.add_heading(level=min(max(block.level, 1), 3))
            _add_docx_runs(paragraph, block.inlines)
        elif block.kind == "paragraph":
            _add_docx_runs(document.add_paragraph(), block.inlines)
        elif block.kind == "list":
            style = "List Number" if block.ordered else "List Bullet"
            for item in block.items:
                _add_docx_runs(document.add_paragraph(style=style), item)
        elif block.kind == "quote":
            for nested in block.blocks:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.left_indent = Inches(0.25)
                _add_docx_runs(paragraph, nested.inlines or [Inline(nested.text)])
        elif block.kind == "code":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(block.text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif block.kind == "table":
            cols = max(len(block.headers), max((len(row) for row in block.rows), default=0))
            if cols == 0:
                continue
            rows = len(block.rows) + (1 if block.headers else 0)
            table = document.add_table(rows=rows, cols=cols)
            table.style = "Table Grid"
            row_offset = 0
            if block.headers:
                for idx, cell in enumerate(block.headers):
                    _add_docx_runs(table.cell(0, idx).paragraphs[0], cell)
                    for run in table.cell(0, idx).paragraphs[0].runs:
                        run.bold = True
                row_offset = 1
            for row_idx, row in enumerate(block.rows):
                for col_idx, cell in enumerate(row):
                    _add_docx_runs(table.cell(row_idx + row_offset, col_idx).paragraphs[0], cell)

    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def _font(size: int):
    from PIL import ImageFont

    candidates = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for path in candidates:
        if os.path.exists(path):
            try:
                return ImageFont.truetype(path, size=size)
            except OSError:
                pass
    return ImageFont.load_default()


def _wrap_text(draw, text: str, font, width: int) -> list[str]:
    lines: list[str] = []
    for raw_line in text.splitlines() or [""]:
        words = raw_line.split(" ")
        current = ""
        for word in words:
            candidate = word if not current else f"{current} {word}"
            if draw.textbbox((0, 0), candidate, font=font)[2] <= width:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = word
        lines.append(current)
    return lines


def render_image(blocks: list[Block], target_format: str) -> bytes:
    from PIL import Image, ImageDraw

    margin = 72
    width = IMAGE_WIDTH
    content_width = width - (margin * 2)
    body_font = _font(24)
    small_font = _font(20)
    code_font = _font(19)
    heading_fonts = {1: _font(42), 2: _font(34), 3: _font(28)}

    scratch = Image.new("RGB", (width, 200), "white")
    draw = ImageDraw.Draw(scratch)
    commands = []
    y = margin

    def add_text(text: str, font, fill=(31, 35, 40), indent=0, spacing=10):
        nonlocal y
        for line in _wrap_text(draw, text, font, content_width - indent):
            commands.append(("text", margin + indent, y, line, font, fill))
            bbox = draw.textbbox((0, 0), line or " ", font=font)
            y += (bbox[3] - bbox[1]) + spacing
        y += spacing

    for block in blocks:
        if block.kind == "heading":
            add_text(_node_text(block.inlines), heading_fonts.get(min(block.level, 3), heading_fonts[3]), fill=(17, 24, 39), spacing=12)
        elif block.kind == "paragraph":
            add_text(_node_text(block.inlines), body_font)
        elif block.kind == "list":
            for idx, item in enumerate(block.items, start=1):
                bullet = f"{idx}. " if block.ordered else "- "
                add_text(bullet + _node_text(item), body_font, indent=18, spacing=8)
        elif block.kind == "quote":
            quote_text = "\n".join(_node_text(nested.inlines) or nested.text for nested in block.blocks)
            add_text(quote_text, body_font, fill=(87, 96, 106), indent=24)
        elif block.kind == "code":
            add_text(block.text, code_font, fill=(36, 41, 47), indent=18, spacing=8)
        elif block.kind == "table":
            table_lines = []
            if block.headers:
                table_lines.append(" | ".join(_node_text(cell) for cell in block.headers))
            for row in block.rows:
                table_lines.append(" | ".join(_node_text(cell) for cell in row))
            add_text("\n".join(table_lines), small_font, fill=(36, 41, 47))

    height = max(y + margin, 400)
    if height > IMAGE_MAX_HEIGHT or (width * height) > IMAGE_MAX_PIXELS:
        raise ValueError("Markdown document is too long to render as one image. Use PDF for very long documents.")

    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    for _kind, x, y_pos, text, font, fill in commands:
        draw.text((x, y_pos), text, font=font, fill=fill)

    out = io.BytesIO()
    fmt = "JPEG" if target_format in ("jpg", "jpeg") else target_format.upper()
    image.save(out, format=fmt, quality=92)
    return out.getvalue()


def convert_markdown(data: bytes, target_format: str) -> bytes:
    target = (target_format or "").lower()
    if target not in TARGET_FORMATS:
        raise ValueError("target_format must be one of: docx, jpeg, jpg, pdf, png, tiff")
    blocks = parse_markdown(_decode_text(data))
    if target == "pdf":
        return render_pdf(blocks)
    if target == "docx":
        return render_docx(blocks)
    return render_image(blocks, target)


def _output_filename(body: dict, file_key: str, target_format: str) -> str:
    original = body.get("file_name") or os.path.basename(file_key) or "markdown.md"
    stem, _ext = os.path.splitext(os.path.basename(original))
    return f"{stem or 'markdown'}.{target_format}"


def handler(event, context):
    body = json.loads(event["Records"][0]["body"])
    job_id = body["job_id"]
    file_key = body["file_key"]
    target_format = (body.get("target_format") or "").lower()

    try:
        dynamo.update_job(job_id, status="PROCESSING")
        data = s3.get_bytes(file_key)
        result = convert_markdown(data, target_format)
        out_key = s3.make_output_key(job_id, file_key, _output_filename(body, file_key, target_format))
        s3.put_bytes(out_key, result)
        dynamo.mark_done(job_id, out_key)
        log.info("markdown_convert done", extra={"job_id": job_id, "target_format": target_format})
    except Exception as exc:
        log.exception("markdown_convert failed: %s", exc)
        dynamo.mark_failed(job_id, str(exc))
        raise

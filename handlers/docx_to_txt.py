# handlers/docx_to_txt.py - Extracts plain text from a .docx file.
#
# python-docx reads the XML inside the docx archive and returns paragraphs
# in order. We join with \n and write as UTF-8. Complex content (tables,
# headers/footers) is NOT captured - those are separate APIs in python-docx
# and a future iteration can add them if users ask.

import io
import html
import json as _json
import os

import dynamo
import s3
from logger import get_logger

log = get_logger(__name__)


def _extract_docx_lines(docx_bytes: bytes) -> list[str]:
    """Read a .docx from bytes and return text lines in document order."""
    # Import inside function to defer the (non-trivial) python-docx import
    # cost until this handler is actually called.
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    lines: list[str] = []
    for paragraph in doc.paragraphs:
        # python-docx exposes the paragraph text directly; empty paragraphs
        # become empty strings, which we keep - they preserve visual spacing.
        lines.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text.strip() for cell in row.cells))
    return lines


def _extract_docx_text(docx_bytes: bytes) -> bytes:
    joined = "\n".join(_extract_docx_lines(docx_bytes))
    return joined.encode("utf-8")


def _extract_docx_markdown(docx_bytes: bytes) -> bytes:
    from docx.enum.style import WD_STYLE_TYPE
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            parts.append("")
            continue
        style = paragraph.style
        style_name = (style.name if style and style.type == WD_STYLE_TYPE.PARAGRAPH else "").lower()
        if style_name.startswith("heading"):
            level_text = "".join(ch for ch in style_name if ch.isdigit())
            level = min(max(int(level_text or "1"), 1), 6)
            parts.append(f"{'#' * level} {text}")
        elif "list bullet" in style_name:
            parts.append(f"- {text}")
        elif "list number" in style_name:
            parts.append(f"1. {text}")
        else:
            parts.append(text)
    for table in doc.tables:
        rows = [[" ".join(cell.text.split()) for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        parts.append("")
        parts.append("| " + " | ".join(rows[0]) + " |")
        parts.append("| " + " | ".join("---" for _ in rows[0]) + " |")
        for row in rows[1:]:
            parts.append("| " + " | ".join(row) + " |")
    return "\n".join(parts).strip().encode("utf-8")


def _extract_docx_html(docx_bytes: bytes) -> bytes:
    parts = ["<!doctype html>", "<html>", "<body>"]
    for line in _extract_docx_lines(docx_bytes):
        # Preserve the document order while normalizing to HTML paragraphs.
        # This keeps the output readable even when the source contained only
        # loosely structured Word content.
        parts.append(f"<p>{html.escape(line)}</p>" if line else "<p>&nbsp;</p>")
    parts.append("</body></html>")
    return "\n".join(parts).encode("utf-8")


def _output_filename(body: dict, file_key: str, target_format: str) -> str:
    original = body.get("file_name") or os.path.basename(file_key) or "document.docx"
    stem, _ext = os.path.splitext(os.path.basename(original))
    return f"{stem or 'document'}.{target_format}"


def handler(event, context):
    body = _json.loads(event["Records"][0]["body"])
    job_id = body["job_id"]
    file_key = body["file_key"]
    operation = body.get("operation", "docx_to_txt")
    params = body.get("params") or {}
    # Default to plain text for docx_to_txt and switch to Markdown or HTML
    # only for the explicit conversion variants in the catalog.
    target_format = (params.get("target_format") or body.get("target_format") or ("md" if operation == "docx_to_md" else ("html" if operation == "docx_to_html" else "txt"))).lower()

    try:
        dynamo.update_job(job_id, status="PROCESSING")
        data = s3.get_bytes(file_key)
        if target_format == "html":
            result = _extract_docx_html(data)
        elif target_format == "md":
            result = _extract_docx_markdown(data)
        elif target_format == "txt":
            result = _extract_docx_text(data)
        else:
            raise ValueError("target_format must be one of: html, md, txt")
        out_key = s3.make_output_key(job_id, file_key, _output_filename(body, file_key, target_format))
        s3.put_bytes(out_key, result)
        dynamo.mark_done(job_id, out_key)
        log.info("docx_to_txt done", extra={"job_id": job_id, "target_format": target_format})
    except Exception as exc:
        log.exception("docx_to_txt failed: %s", exc)
        dynamo.mark_failed(job_id, str(exc))
        raise

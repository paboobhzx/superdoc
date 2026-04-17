import io
import json

import dynamo
import s3
from logger import get_logger

log = get_logger(__name__)


def _apply_ops_docx(data: bytes, ops: list) -> bytes:
    from docx import Document
    doc = Document(io.BytesIO(data))
    for op in ops:
        action = op.get("action")
        if action == "replace_text":
            find = op.get("find", "")
            replace = op.get("replace", "")
            for para in doc.paragraphs:
                if find in para.text:
                    for run in para.runs:
                        run.text = run.text.replace(find, replace)
        elif action == "add_paragraph":
            doc.add_paragraph(op.get("text", ""))
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _apply_ops_xlsx(data: bytes, ops: list) -> bytes:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data))
    ws = wb.active
    for op in ops:
        action = op.get("action")
        if action == "set_cell":
            row = op.get("row", 1)
            col = op.get("col", 1)
            value = op.get("value", "")
            ws.cell(row=row, column=col, value=value)
    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


def _process(data: bytes, body: dict) -> tuple[bytes, str]:
    file_name = body.get("file_name", "file.docx").lower()
    ops = body.get("ops", [])
    if file_name.endswith(".xlsx"):
        return _apply_ops_xlsx(data, ops), "xlsx"
    return _apply_ops_docx(data, ops), "docx"


def handler(event, context):
    body = json.loads(event["Records"][0]["body"])
    job_id = body["job_id"]
    file_key = body["file_key"]
    try:
        dynamo.update_job(job_id, status="PROCESSING")
        data = s3.get_bytes(file_key)
        result, ext = _process(data, body)
        out_key = s3.make_output_key(job_id, file_key, f"edited.{ext}")
        s3.put_bytes(out_key, result)
        dynamo.mark_done(job_id, out_key)
        log.info("doc_edit done", extra={"job_id": job_id})
    except Exception as exc:
        log.exception("doc_edit failed: %s", exc)
        dynamo.mark_failed(job_id, str(exc))
        raise

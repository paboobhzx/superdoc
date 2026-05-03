import html
import io
import json as _json
import os
import shutil
import subprocess
import tempfile
import zipfile

import dynamo
import s3
from logger import get_logger

log = get_logger(__name__)

_LIBREOFFICE_BIN = os.environ.get("LIBREOFFICE_BIN", "libreoffice")


def _load_sheet_rows(xlsx_bytes: bytes, sheet_name: str | None):
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(xlsx_bytes), read_only=True, data_only=True)
    try:
        if sheet_name is None:
            ws = wb.worksheets[0]
        else:
            if sheet_name not in wb.sheetnames:
                raise ValueError(f"sheet not found: {sheet_name}")
            ws = wb[sheet_name]
        return [list(row) for row in ws.iter_rows(values_only=True)]
    finally:
        wb.close()


def _normalize_rows(rows):
    normalized: list[list[str]] = []
    for row in rows:
        current: list[str] = []
        for cell in row:
            current.append("" if cell is None else str(cell))
        normalized.append(current)
    return normalized


def _rows_to_docx(rows) -> bytes:
    from docx import Document
    from docx.shared import Pt

    normalized = _normalize_rows(rows)
    document = Document()
    document.styles["Normal"].font.size = Pt(10.5)
    if not normalized:
        document.add_paragraph("")
    else:
        width = max(len(row) for row in normalized)
        padded = [row + [""] * (width - len(row)) for row in normalized]
        table = document.add_table(rows=len(padded), cols=width)
        table.style = "Table Grid"
        for row_idx, row in enumerate(padded):
            for col_idx, cell in enumerate(row):
                table.cell(row_idx, col_idx).text = cell
    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def _xlsx_to_pdf(xlsx_bytes: bytes, sheet_name: str | None = None) -> bytes:
    from openpyxl import load_workbook

    if sheet_name:
        workbook_path = _select_sheet_workbook(xlsx_bytes, sheet_name)
        try:
            with open(workbook_path, "rb") as fh:
                return _office_to_pdf(fh.read(), "input.xlsx")
        finally:
            try:
                os.unlink(workbook_path)
            except FileNotFoundError:
                pass
    return _office_to_pdf(xlsx_bytes, "input.xlsx")


def _xlsx_to_image_zip(xlsx_bytes: bytes, sheet_name: str | None, target_format: str = "png", dpi: int = 150) -> bytes:
    pdf_bytes = _xlsx_to_pdf(xlsx_bytes, sheet_name=sheet_name)
    return _render_pdf_to_zip(pdf_bytes, target_format=target_format, dpi=dpi)


def _xlsx_to_docx(xlsx_bytes: bytes, sheet_name: str | None = None) -> bytes:
    rows = _load_sheet_rows(xlsx_bytes, sheet_name)
    return _rows_to_docx(rows)


def _xlsx_to_html(xlsx_bytes: bytes, sheet_name: str | None = None) -> bytes:
    rows = _load_sheet_rows(xlsx_bytes, sheet_name)
    normalized = _normalize_rows(rows)
    parts = ["<!doctype html>", "<html>", "<body>", "<table>"]
    if normalized:
        parts.append("<thead><tr>")
        for cell in normalized[0]:
            parts.append(f"<th>{html.escape(cell)}</th>")
        parts.append("</tr></thead>")
        if len(normalized) > 1:
            parts.append("<tbody>")
            for row in normalized[1:]:
                parts.append("<tr>")
                for cell in row:
                    parts.append(f"<td>{html.escape(cell)}</td>")
                parts.append("</tr>")
            parts.append("</tbody>")
    parts.extend(["</table>", "</body>", "</html>"])
    return "\n".join(parts).encode("utf-8")


def _render_pdf_to_zip(pdf_bytes: bytes, target_format: str, dpi: int) -> bytes:
    import pymupdf

    target = (target_format or "png").lower()
    if target == "jpeg":
        target = "jpg"
    if target not in {"png", "jpg"}:
        raise ValueError("target_format must be one of: jpg, png")

    doc = pymupdf.open(stream=pdf_bytes, filetype="pdf")
    try:
        buffer = io.BytesIO()
        matrix = pymupdf.Matrix(dpi / 72.0, dpi / 72.0)
        with zipfile.ZipFile(buffer, mode="w", compression=zipfile.ZIP_STORED) as zf:
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                pix = page.get_pixmap(matrix=matrix, alpha=False)
                image_bytes = pix.tobytes("jpeg" if target == "jpg" else "png")
                zf.writestr(f"page-{page_num + 1:04d}.{target}", image_bytes)
        return buffer.getvalue()
    finally:
        doc.close()


def _select_sheet_workbook(xlsx_bytes: bytes, sheet_name: str) -> str:
    import io

    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(xlsx_bytes))
    path = ""
    try:
        if sheet_name:
            if sheet_name not in wb.sheetnames:
                raise ValueError(f"sheet not found: {sheet_name}")
            ws = wb[sheet_name]
        else:
            ws = wb.worksheets[0]
        selected_title = ws.title
        wb.active = wb.index(ws)
        for worksheet in wb.worksheets:
            worksheet.sheet_state = "visible" if worksheet.title == selected_title else "hidden"
        fd, path = tempfile.mkstemp(prefix="selected-sheet-", suffix=".xlsx")
        os.close(fd)
        wb.save(path)
        return path
    finally:
        wb.close()


def _office_to_pdf(source_bytes: bytes, source_name: str) -> bytes:
    if not shutil.which(_LIBREOFFICE_BIN):
        raise RuntimeError(f"LibreOffice executable not found: {_LIBREOFFICE_BIN}")

    with tempfile.TemporaryDirectory(prefix="office-to-pdf-") as workdir:
        outdir = os.path.join(workdir, "out")
        profile = os.path.join(workdir, "profile")
        os.makedirs(outdir, exist_ok=True)
        os.makedirs(profile, exist_ok=True)

        source_path = os.path.join(workdir, source_name)
        with open(source_path, "wb") as fh:
            fh.write(source_bytes)

        cmd = [
            _LIBREOFFICE_BIN,
            "--headless",
            "--nologo",
            "--nodefault",
            "--nofirststartwizard",
            "--nolockcheck",
            f"-env:UserInstallation=file://{profile}",
            "--convert-to",
            "pdf",
            "--outdir",
            outdir,
            source_path,
        ]
        completed = subprocess.run(
            cmd,
            check=False,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=int(os.environ.get("LIBREOFFICE_TIMEOUT_SECONDS", "240")),
        )
        if completed.returncode != 0:
            stderr = completed.stderr.decode("utf-8", errors="replace").strip()
            stdout = completed.stdout.decode("utf-8", errors="replace").strip()
            detail = stderr or stdout or f"exit code {completed.returncode}"
            raise RuntimeError(f"LibreOffice PDF export failed: {detail}")

        pdf_path = os.path.join(outdir, f"{os.path.splitext(source_name)[0]}.pdf")
        if not os.path.exists(pdf_path):
            raise RuntimeError("LibreOffice PDF export did not produce an output file")

        with open(pdf_path, "rb") as fh:
            result = fh.read()
        if not result.startswith(b"%PDF"):
            raise RuntimeError("LibreOffice PDF export produced an invalid PDF")
        return result


def handler(event, context):
    body = _json.loads(event["Records"][0]["body"])
    job_id = body["job_id"]
    file_key = body["file_key"]
    operation = body.get("operation", "xlsx_to_pdf")
    params = body.get("params") or {}
    target_format = (params.get("target_format") or body.get("target_format") or "pdf").lower()
    sheet_name = params.get("sheet") or body.get("sheet")
    dpi = int(params.get("dpi") or body.get("dpi") or 150)

    try:
        dynamo.update_job(job_id, status="PROCESSING")
        data = s3.get_bytes(file_key)
        if target_format == "docx" or operation == "xlsx_to_docx":
            result = _xlsx_to_docx(data, sheet_name=sheet_name)
            output_name = "output.docx"
        elif target_format in ("png", "jpg", "jpeg") or operation == "xlsx_to_image":
            result = _xlsx_to_image_zip(data, sheet_name=sheet_name, target_format=target_format, dpi=dpi)
            output_name = "pages.zip"
        elif target_format == "html" or operation == "xlsx_to_html":
            result = _xlsx_to_html(data, sheet_name=sheet_name)
            output_name = "output.html"
        else:
            result = _xlsx_to_pdf(data, sheet_name=sheet_name)
            output_name = "output.pdf"
        out_key = s3.make_output_key(job_id, file_key, output_name)
        s3.put_bytes(out_key, result)
        dynamo.mark_done(job_id, out_key)
        log.info("xlsx_to_pdf done", extra={"job_id": job_id, "target_format": target_format, "operation": operation})
    except Exception as exc:
        log.exception("xlsx_to_pdf failed: %s", exc)
        dynamo.mark_failed(job_id, str(exc))
        raise

import copy
import io
import json
import os
import shutil
import subprocess
import tempfile

import dynamo
import output_naming
import s3
from logger import get_logger

log = get_logger(__name__)

_LIBREOFFICE_BIN = os.environ.get("LIBREOFFICE_BIN", "libreoffice")
_QA_THRESHOLD = float(os.environ.get("PDF_TO_DOCX_QA_THRESHOLD", "0.7"))
_CHUNK_SIZE = 40  # pages per chunk for large document splits


def _pdf_to_docx_libreoffice(pdf_bytes: bytes) -> bytes:
    if not shutil.which(_LIBREOFFICE_BIN):
        raise RuntimeError(f"LibreOffice executable not found: {_LIBREOFFICE_BIN}")

    with tempfile.TemporaryDirectory(prefix="pdf-to-docx-") as workdir:
        outdir = os.path.join(workdir, "out")
        profile = os.path.join(workdir, "profile")
        os.makedirs(outdir, exist_ok=True)
        os.makedirs(profile, exist_ok=True)

        source_path = os.path.join(workdir, "input.pdf")
        with open(source_path, "wb") as fh:
            fh.write(pdf_bytes)

        cmd = [
            _LIBREOFFICE_BIN,
            "--headless",
            "--nologo",
            "--nodefault",
            "--nofirststartwizard",
            "--nolockcheck",
            f"-env:UserInstallation=file://{profile}",
            "--infilter=writer_pdf_import",
            "--convert-to",
            "docx:MS Word 2007 XML",
            "--outdir",
            outdir,
            source_path,
        ]
        completed = subprocess.run(
            cmd,
            check=False,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=int(os.environ.get("LIBREOFFICE_TIMEOUT_SECONDS", "240")),
        )
        if completed.returncode != 0:
            stderr = completed.stderr.decode("utf-8", errors="replace").strip()
            stdout = completed.stdout.decode("utf-8", errors="replace").strip()
            detail = stderr or stdout or f"exit code {completed.returncode}"
            raise RuntimeError(f"LibreOffice DOCX export failed: {detail}")

        docx_path = os.path.join(outdir, "input.docx")
        if not os.path.exists(docx_path):
            raise RuntimeError("LibreOffice DOCX export did not produce an output file")

        with open(docx_path, "rb") as fh:
            return fh.read()


def _extract_text_docx(pdf_bytes: bytes) -> bytes:
    from docx import Document
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(pdf_bytes))
    doc = Document()
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            doc.add_paragraph(text)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _parse_page_range(page_range: str, total: int) -> list[int]:
    """Parse '1-5,8,10-12' into a sorted 0-indexed list, clamped to [0, total)."""
    if not page_range or not page_range.strip():
        return list(range(total))
    indices: set[int] = set()
    for part in page_range.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            a, _, b = part.partition("-")
            start = max(0, int(a.strip()) - 1)
            end = min(int(b.strip()) - 1, total - 1)
            if start <= end:
                indices.update(range(start, end + 1))
        else:
            idx = int(part.strip()) - 1
            if 0 <= idx < total:
                indices.add(idx)
    return sorted(indices)


def _merge_docx(docx_bytes_list: list[bytes]) -> bytes:
    """Merge DOCX chunks into one document.

    Uses lxml addprevious() to insert all chunk content BEFORE the master's
    sectPr, which must remain the final child of <w:body> per OOXML spec.
    Using body.append() would place content after sectPr and produce the
    'unreadable content' error in Word.
    """
    from docx import Document
    from docx.oxml.ns import qn

    if len(docx_bytes_list) == 1:
        return docx_bytes_list[0]

    _IMAGE_RTYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
    _R_EMBED = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
    _SECT_PR = qn("w:sectPr")

    master = Document(io.BytesIO(docx_bytes_list[0]))

    for chunk_bytes in docx_bytes_list[1:]:
        chunk = Document(io.BytesIO(chunk_bytes))

        children = list(chunk.element.body)
        if children and children[-1].tag == _SECT_PR:
            children = children[:-1]

        if not children:
            continue

        # Add a page break before the new chunk's content.
        # Document.add_page_break() inserts before sectPr via python-docx internals.
        master.add_page_break()

        # Locate the master sectPr — chunk content must be inserted BEFORE it,
        # never appended after it.
        master_sect_pr = master.element.body.find(_SECT_PR)

        for elem in children:
            clone = copy.deepcopy(elem)
            # Remap embedded image relationship IDs to the master document
            for blip in clone.findall(".//" + qn("a:blip")):
                old_rid = blip.get(_R_EMBED)
                if old_rid and old_rid in chunk.part.rels:
                    image_part = chunk.part.rels[old_rid].target_part
                    new_rid = master.part.relate_to(image_part, _IMAGE_RTYPE)
                    blip.set(_R_EMBED, new_rid)
            if master_sect_pr is not None:
                master_sect_pr.addprevious(clone)  # insert before sectPr
            else:
                master.element.body.append(clone)

    out = io.BytesIO()
    master.save(out)
    return out.getvalue()


def _render_page_as_image_docx(page_bytes: bytes) -> bytes:
    """Render a single-page PDF as a full-page image embedded in a DOCX."""
    import pymupdf
    from docx import Document
    from docx.shared import Inches

    usable_w = Inches(7.27)
    usable_h = Inches(10.69)
    zoom = 150 / 72.0
    matrix = pymupdf.Matrix(zoom, zoom)

    doc = pymupdf.open(stream=page_bytes, filetype="pdf")
    try:
        docx_doc = Document()
        section = docx_doc.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

        page = doc.load_page(0)
        pix = page.get_pixmap(matrix=matrix, alpha=False)
        pix_w, pix_h = pix.width, pix.height
        png_bytes = pix.tobytes("png")
        pix = None  # explicit release for memory safety

        img_stream = io.BytesIO(png_bytes)
        png_bytes = None

        if pix_h > 0:
            aspect = pix_w / pix_h
            height_if_fit_w = usable_w / aspect
            if height_if_fit_w > usable_h:
                docx_doc.add_picture(img_stream, height=usable_h)
            else:
                docx_doc.add_picture(img_stream, width=usable_w)
        else:
            docx_doc.add_picture(img_stream, width=usable_w)

        out = io.BytesIO()
        docx_doc.save(out)
        return out.getvalue()
    finally:
        doc.close()


def _count_words_in_text(text: str) -> int:
    return len(text.split())


def _count_words_docx(docx_bytes: bytes) -> int:
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    total = 0
    for para in doc.paragraphs:
        total += _count_words_in_text(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total += _count_words_in_text(cell.text)
    return total


def _pdf2docx_convert(pdf_bytes: bytes, body: dict) -> bytes:
    """Convert a PDF to DOCX using pdf2docx → LibreOffice → text fallback."""
    try:
        from pdf2docx import Converter

        with tempfile.TemporaryDirectory(prefix="pdf2docx-") as workdir:
            src = os.path.join(workdir, "input.pdf")
            stem, _ = os.path.splitext(os.path.basename(body.get("file_name") or "doc.pdf"))
            dst = os.path.join(workdir, f"{stem or 'document'}.docx")
            with open(src, "wb") as fh:
                fh.write(pdf_bytes)
            cv = Converter(src)
            try:
                cv.convert(dst)
            finally:
                cv.close()
            with open(dst, "rb") as fh:
                return fh.read()
    except Exception:
        pass

    try:
        return _pdf_to_docx_libreoffice(pdf_bytes)
    except Exception:
        return _extract_text_docx(pdf_bytes)


def _build_text_fallback_docx(pdf_bytes: bytes) -> bytes:
    """Build a plain-text DOCX from the full PDF using PyMuPDF."""
    import pymupdf
    from docx import Document

    docx_doc = Document()
    doc = pymupdf.open(stream=pdf_bytes, filetype="pdf")
    try:
        for page_num in range(doc.page_count):
            if page_num > 0:
                docx_doc.add_page_break()
            page = doc.load_page(page_num)
            text = page.get_text("text") or ""
            for line in text.split("\n"):
                if line.strip():
                    docx_doc.add_paragraph(line.strip())
    finally:
        doc.close()
    out = io.BytesIO()
    docx_doc.save(out)
    return out.getvalue()


def _build_image_fallback_docx(pdf_bytes: bytes) -> bytes:
    """Render every page of the PDF as a full-page image in a DOCX."""
    import pymupdf
    from docx import Document
    from docx.shared import Inches

    usable_w = Inches(7.27)
    usable_h = Inches(10.69)
    zoom = 150 / 72.0
    matrix = pymupdf.Matrix(zoom, zoom)

    doc = pymupdf.open(stream=pdf_bytes, filetype="pdf")
    try:
        docx_doc = Document()
        section = docx_doc.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

        for page_num in range(doc.page_count):
            if page_num > 0:
                docx_doc.add_page_break()
            page = doc.load_page(page_num)
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            pix_w, pix_h = pix.width, pix.height
            png_bytes = pix.tobytes("png")
            pix = None

            img_stream = io.BytesIO(png_bytes)
            png_bytes = None

            if pix_h > 0:
                aspect = pix_w / pix_h
                height_if_fit_w = usable_w / aspect
                if height_if_fit_w > usable_h:
                    docx_doc.add_picture(img_stream, height=usable_h)
                else:
                    docx_doc.add_picture(img_stream, width=usable_w)
            else:
                docx_doc.add_picture(img_stream, width=usable_w)

        out = io.BytesIO()
        docx_doc.save(out)
        return out.getvalue()
    finally:
        doc.close()


def _single_page_pdf_bytes(reader, page_index: int) -> bytes:
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_page(reader.pages[page_index])
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def _count_words_pdf_page(page_bytes: bytes) -> int:
    try:
        import pymupdf

        doc = pymupdf.open(stream=page_bytes, filetype="pdf")
        try:
            page = doc.load_page(0)
            return _count_words_in_text(page.get_text("text") or "")
        finally:
            doc.close()
    except Exception:
        return 0


def _hybrid_high_fidelity_docx(pdf_bytes: bytes, body: dict) -> bytes:
    """Maximum-effort PDF->DOCX: reconstruct per page, image-fallback weak pages."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(pdf_bytes))
    page_count = len(reader.pages)
    chunks: list[bytes] = []
    used_image_pages = 0
    used_reconstructed_pages = 0
    job_id = body.get("job_id", "unknown")

    for page_index in range(page_count):
        page_bytes = _single_page_pdf_bytes(reader, page_index)
        source_words = _count_words_pdf_page(page_bytes)
        use_image = source_words == 0
        candidate = b""

        if not use_image:
            try:
                candidate = _pdf2docx_convert(page_bytes, body)
                candidate_words = _count_words_docx(candidate)
                ratio = candidate_words / max(source_words, 1)
                use_image = ratio < _QA_THRESHOLD
                log.info(
                    "high_fidelity page qa ratio=%.3f",
                    ratio,
                    extra={"job_id": job_id, "page": page_index + 1},
                )
            except Exception:
                log.exception(
                    "high_fidelity page reconstruction failed",
                    extra={"job_id": job_id, "page": page_index + 1},
                )
                use_image = True

        if use_image:
            chunks.append(_render_page_as_image_docx(page_bytes))
            used_image_pages += 1
        else:
            chunks.append(candidate)
            used_reconstructed_pages += 1

    log.info(
        "high_fidelity hybrid result",
        extra={
            "job_id": job_id,
            "pages": page_count,
            "reconstructed_pages": used_reconstructed_pages,
            "image_pages": used_image_pages,
        },
    )
    return _merge_docx(chunks)


def _process(pdf_bytes: bytes, body: dict) -> bytes:
    from pypdf import PdfReader, PdfWriter

    params = body.get("params") if isinstance(body.get("params"), dict) else {}
    if not params:
        params = {
            key: body[key]
            for key in ("page_range", "high_fidelity", "fallback_strategy")
            if key in body
        }
    page_range_str = (params.get("page_range") or "").strip()
    high_fidelity = bool(params.get("high_fidelity", False))
    fallback_strategy = (params.get("fallback_strategy") or "text").lower()
    job_id = body.get("job_id", "unknown")

    reader = PdfReader(io.BytesIO(pdf_bytes))
    total_pages = len(reader.pages)
    page_indices = _parse_page_range(page_range_str, total_pages)
    n_pages = len(page_indices)

    log.info(
        "pdf_to_docx conversion planned",
        extra={
            "job_id": job_id,
            "total_pages": total_pages,
            "pages_to_convert": n_pages,
            "high_fidelity": high_fidelity,
            "fallback_strategy": fallback_strategy,
        },
    )

    # Extract page subset if a range was specified
    if page_range_str and n_pages < total_pages:
        writer = PdfWriter()
        for idx in page_indices:
            writer.add_page(reader.pages[idx])
        buf = io.BytesIO()
        writer.write(buf)
        pdf_bytes = buf.getvalue()

    # ── Mode 1: hybrid high-fidelity reconstruction ──────────────────────────
    if high_fidelity:
        log.info("mode1 (hybrid high-fidelity)", extra={"job_id": job_id})
        try:
            dynamo.record_page_result(job_id, page=0, mode_used="hybrid")
        except Exception:
            pass
        return _hybrid_high_fidelity_docx(pdf_bytes, body)

    # ── Mode 0: pdf2docx on the full document, QA gate, fallback ─────────────
    #
    # We run pdf2docx on the whole document (not per-page) because:
    # 1. pdf2docx is designed for full documents — it resolves cross-page layout
    # 2. Per-page extraction + merge produced "unreadable content" errors in Word
    # 3. The QA gate can still detect catastrophic failures and fall back
    #
    # For large PDFs that exceed CHUNK_SIZE, split by page range, convert each
    # chunk as a whole, then merge. This preserves whole-chunk layout continuity.

    # Count total extractable words for QA gate
    total_pdf_words = 0
    try:
        import pymupdf
        fitz_doc = pymupdf.open(stream=pdf_bytes, filetype="pdf")
        try:
            for page_num in range(fitz_doc.page_count):
                page = fitz_doc.load_page(page_num)
                total_pdf_words += _count_words_in_text(page.get_text("text") or "")
        finally:
            fitz_doc.close()
    except Exception:
        pass

    if n_pages <= _CHUNK_SIZE:
        docx_result = _pdf2docx_convert(pdf_bytes, body)
    else:
        # Split into sequential chunks, convert each, then merge
        docx_chunks: list[bytes] = []
        reader2 = PdfReader(io.BytesIO(pdf_bytes))
        for chunk_start in range(0, n_pages, _CHUNK_SIZE):
            chunk_end = min(chunk_start + _CHUNK_SIZE, n_pages)
            chunk_writer = PdfWriter()
            for i in range(chunk_start, chunk_end):
                chunk_writer.add_page(reader2.pages[i])
            chunk_buf = io.BytesIO()
            chunk_writer.write(chunk_buf)
            log.info(
                "pdf_to_docx chunk",
                extra={"job_id": job_id, "chunk_start": chunk_start + 1, "chunk_end": chunk_end},
            )
            docx_chunks.append(_pdf2docx_convert(chunk_buf.getvalue(), body))
        docx_result = _merge_docx(docx_chunks)

    # QA gate: check word retention. If pdf has no text (scanned), skip.
    if total_pdf_words > 0:
        try:
            docx_words = _count_words_docx(docx_result)
            ratio = docx_words / total_pdf_words
            log.info(
                "qa_gate ratio=%.3f (%d pdf / %d docx)",
                ratio, total_pdf_words, docx_words,
                extra={"job_id": job_id},
            )
            if ratio < _QA_THRESHOLD:
                log.warning(
                    "qa_gate triggered ratio=%.3f < %.2f, fallback=%s",
                    ratio, _QA_THRESHOLD, fallback_strategy,
                    extra={"job_id": job_id},
                )
                if fallback_strategy == "image":
                    docx_result = _build_image_fallback_docx(pdf_bytes)
                else:
                    docx_result = _build_text_fallback_docx(pdf_bytes)
        except Exception:
            log.exception("qa_gate error, returning pdf2docx result", extra={"job_id": job_id})

    try:
        mode = "pdf2docx"
        dynamo.record_page_result(job_id, page=0, mode_used=mode)
    except Exception:
        pass

    return docx_result


def _output_filename(body: dict, file_key: str) -> str:
    return output_naming.output_filename("pdf_to_docx", body, file_key, "docx")


def handler(event, context):
    if event.get("_warmup"):
        log.info("warmup ping received")
        return
    body = json.loads(event["Records"][0]["body"])
    job_id = body["job_id"]
    file_key = body["file_key"]
    try:
        dynamo.update_job(job_id, status="PROCESSING")
        data = s3.get_bytes(file_key)
        result = _process(data, body)
        out_key = s3.make_output_key(job_id, file_key, _output_filename(body, file_key))
        s3.put_bytes(out_key, result)
        dynamo.mark_done(job_id, out_key)
        log.info("pdf_to_docx done", extra={"job_id": job_id})
    except Exception as exc:
        log.exception("pdf_to_docx failed: %s", exc)
        dynamo.mark_failed(job_id, str(exc))
        raise

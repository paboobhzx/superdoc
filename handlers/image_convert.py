import io
import json
import os

import dynamo
import ocr
import s3
from logger import get_logger
from PIL import Image

log = get_logger(__name__)

_FORMAT_MAP = {
    "jpg": "JPEG",
    "jpeg": "JPEG",
    "png": "PNG",
    "webp": "WEBP",
    "gif": "GIF",
    "bmp": "BMP",
    "tiff": "TIFF",
}

_DOCUMENT_FORMATS = {"txt", "md", "docx"}


def _process(data: bytes, body: dict) -> tuple[bytes, str]:
    params = body.get("params") or {}
    # The image worker is shared by format conversion and OCR-backed document
    # exports. Document targets bypass PIL entirely because the OCR path needs
    # the extracted lines, not a raster re-encode.
    target_format = (params.get("target_format") or body.get("target_format") or "png").lower()
    if target_format in _DOCUMENT_FORMATS:
        return _ocr_to_document(data, target_format), target_format
    if target_format not in _FORMAT_MAP:
        raise ValueError("target_format must be one of: docx, gif, jpg, jpeg, md, png, txt, webp")
    pil_format = _FORMAT_MAP.get(target_format, "PNG")

    img = Image.open(io.BytesIO(data))
    if pil_format == "JPEG" and img.mode in ("RGBA", "P"):
        img = img.convert("RGB")

    out = io.BytesIO()
    img.save(out, format=pil_format)
    return out.getvalue(), target_format


def _ocr_to_document(data: bytes, target_format: str) -> bytes:
    lines = ocr.extract_lines(_normalize_for_ocr(data))
    text = "\n".join(lines).strip()
    if target_format == "txt":
        return (text + ("\n" if text else "")).encode("utf-8")
    if target_format == "md":
        return _lines_to_markdown(lines).encode("utf-8")
    if target_format == "docx":
        return _lines_to_docx(lines)
    raise ValueError("target_format must be one of: docx, md, txt")


def _normalize_for_ocr(data: bytes) -> bytes:
    img = Image.open(io.BytesIO(data))
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    out = io.BytesIO()
    img.save(out, format="PNG")
    return out.getvalue()


def _lines_to_markdown(lines: list[str]) -> str:
    return "\n\n".join(line.strip() for line in lines if line.strip())


def _lines_to_docx(lines: list[str]) -> bytes:
    from docx import Document

    document = Document()
    if lines:
        for line in lines:
            document.add_paragraph(line)
    else:
        document.add_paragraph("")
    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def handler(event, context):
    body = json.loads(event["Records"][0]["body"])
    job_id = body["job_id"]
    file_key = body["file_key"]
    try:
        dynamo.update_job(job_id, status="PROCESSING")
        data = s3.get_bytes(file_key)
        result, ext = _process(data, body)
        out_key = s3.make_output_key(job_id, file_key, f"converted.{ext}")
        s3.put_bytes(out_key, result)
        dynamo.mark_done(job_id, out_key)
        log.info("image_convert done", extra={"job_id": job_id})
    except Exception as exc:
        log.exception("image_convert failed: %s", exc)
        dynamo.mark_failed(job_id, str(exc))
        raise
